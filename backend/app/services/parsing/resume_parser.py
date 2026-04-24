"""
app/services/parsing/resume_parser.py

Stage 2 — Resume Parsing Layer

Responsibilities:
  1. Extract raw text from PDF / DOCX
  2. Detect layout type (single vs multi-column)
  3. Detect and slice sections (Experience, Skills, Education, Projects, ...)
  4. Score parse quality (0.0 – 1.0)

Output: ParsedResume dataclass
"""
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

# ── Section header vocabulary ─────────────────────────────────────────────────
SECTION_ALIASES: dict[str, str] = {
    # Experience
    "work experience": "experience",
    "experience": "experience",
    "professional experience": "experience",
    "employment history": "experience",
    "work history": "experience",
    "career history": "experience",
    "professional background": "experience",
    "positions held": "experience",
    "relevant experience": "experience",
    "internship experience": "experience",
    "internships": "experience",

    # Skills
    "skills": "skills",
    "technical skills": "skills",
    "core competencies": "skills",
    "competencies": "skills",
    "technologies": "skills",
    "tech stack": "skills",
    "tools & technologies": "skills",
    "tools and technologies": "skills",
    "key skills": "skills",
    "areas of expertise": "skills",
    "expertise": "skills",
    "proficiencies": "skills",
    "programming languages": "skills",
    "languages & frameworks": "skills",

    # Education
    "education": "education",
    "academic background": "education",
    "qualifications": "education",
    "academic qualifications": "education",
    "educational background": "education",
    "degrees": "education",
    "academics": "education",

    # Projects
    "projects": "projects",
    "personal projects": "projects",
    "side projects": "projects",
    "portfolio": "projects",
    "academic projects": "projects",
    "notable projects": "projects",
    "key projects": "projects",
    "open source": "projects",
    "open source contributions": "projects",

    # Certifications
    "certifications": "certifications",
    "certifications and achievements": "certifications",
    "certifications & achievements": "certifications",
    "certificates": "certifications",
    "licenses": "certifications",
    "credentials": "certifications",
    "professional certifications": "certifications",
    "achievements": "certifications",
    "accomplishments": "certifications",
    "honors": "certifications",
    "awards": "certifications",
    "honors & awards": "certifications",

    # Summary
    "summary": "summary",
    "profile": "summary",
    "professional summary": "summary",
    "career objective": "summary",
    "objective": "summary",
    "about me": "summary",
    "overview": "summary",
    "executive summary": "summary",

    # Languages
    "languages": "languages",
    "spoken languages": "languages",

    # Extracurricular — map to misc but recognised so no noise
    "extracurricular activities": "extracurricular",
    "extracurricular": "extracurricular",
    "activities": "extracurricular",
    "volunteer": "extracurricular",
    "volunteering": "extracurricular",
    # NOTE: keep "interests"/"hobbies" as misc sections (tests expect misc preservation)
}

CANONICAL_SECTIONS = {
    "experience", "skills", "education", "projects",
    "certifications", "summary", "languages", "extracurricular",
}

# Sections to skip during skill extraction — not useful for skills
SKIP_FOR_SKILLS = {"extracurricular", "languages"}

# Minimum chars for a section to be considered non-empty
MIN_SECTION_LENGTH = 5

# Lines that are never section headers regardless of formatting
IGNORE_LINES = {
    "references available upon request",
    "references",
    "declaration",
    "i hereby declare",
    "place",
    "date",
}


@dataclass
class ParsedResume:
    raw_text: str = ""
    layout_type: str = "single"
    sections: dict[str, str] = field(default_factory=dict)
    misc_sections: dict[str, str] = field(default_factory=dict)
    parse_confidence: float = 0.0
    warnings: list[str] = field(default_factory=list)


class ResumeParser:

    def parse(self, file_path: str) -> ParsedResume:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            raw_text, layout_type = self._extract_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            raw_text, layout_type = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        result = ParsedResume(raw_text=raw_text, layout_type=layout_type)

        if not raw_text.strip():
            result.warnings.append("Empty text extracted — file may be image-based or corrupt")
            result.parse_confidence = 0.0
            return result

        result.sections, result.misc_sections = self._detect_sections(raw_text)
        result.parse_confidence = self._score_confidence(result)
        return result

    def _extract_pdf(self, file_path: str) -> tuple[str, str]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not installed — run: pip install pdfplumber")

        pages_text = []
        layout_type = "single"
        column_votes = 0
        total_pages = 0

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                words = page.extract_words()
                if words:
                    page_mid = page.width / 2
                    left_words  = [w for w in words if float(w["x0"]) < page_mid - 20]
                    right_words = [w for w in words if float(w["x0"]) > page_mid + 20]
                    if len(left_words) > 10 and len(right_words) > 10:
                        ratio = min(len(left_words), len(right_words)) / max(len(left_words), len(right_words))
                        if ratio > 0.3:
                            column_votes += 1

                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                pages_text.append(text)

        if total_pages > 0 and column_votes / total_pages > 0.4:
            layout_type = "multi_column"

        return "\n\n".join(pages_text), layout_type

    def _extract_docx(self, file_path: str) -> tuple[str, str]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed — run: pip install python-docx")

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        layout_type = "single"
        for table in doc.tables:
            if len(table.columns) >= 2:
                layout_type = "multi_column"
                break

        return "\n".join(paragraphs), layout_type

    def _detect_sections(self, text: str) -> tuple[dict[str, str], dict[str, str]]:
        lines = text.split("\n")
        header_positions: list[tuple[int, str, bool]] = []

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
            canonical, recognised = self._classify_header(stripped)
            if canonical:
                header_positions.append((i, canonical, recognised))

        if not header_positions:
            return {}, {"full_text": text}

        sections: dict[str, str] = {}
        misc_sections: dict[str, str] = {}

        for idx, (line_i, name, recognised) in enumerate(header_positions):
            start = line_i + 1
            end   = header_positions[idx + 1][0] if idx + 1 < len(header_positions) else len(lines)
            content = "\n".join(lines[start:end]).strip()

            if len(content) < MIN_SECTION_LENGTH:
                continue

            if recognised:
                if name in sections:
                    sections[name] += "\n" + content
                else:
                    sections[name] = content
            else:
                misc_sections[name] = content

        return sections, misc_sections

    def _classify_header(self, line: str) -> tuple[Optional[str], bool]:
        # Too long or too short
        if len(line) > 60 or len(line) < 4:
            return None, False

        # Ends with period — likely a sentence fragment not a header
        if line.rstrip().endswith(".") or line.rstrip().endswith(","):
            return None, False

        # Heavy digit content — percentage, date, score
        digit_ratio = sum(c.isdigit() for c in line) / max(len(line), 1)
        if digit_ratio > 0.3:
            return None, False

        # Clean for lookup — defined early so all checks below can use it
        cleaned = line.rstrip(":").strip().lower()

        # Explicitly ignored lines
        if cleaned in IGNORE_LINES:
            return None, False

        # Extracurricular / non-skill sections — recognised but excluded from misc
        ignore_as_misc = {
            "extracurricular activities", "extracurricular",
            "references", "declaration",
            "personal information", "personal details",
        }

        # Person's name check — likely a name, not a header
        words = line.split()
        if len(words) >= 3 and all(w[0].isupper() for w in words if w):
            if cleaned not in SECTION_ALIASES:
                return None, False
        if len(words) == 2 and line.istitle() and cleaned not in SECTION_ALIASES:
            return None, False

        # Must be mostly alphabetic
        alpha_ratio = sum(c.isalpha() or c in " &/-" for c in line) / max(len(line), 1)
        if alpha_ratio < 0.6:
            return None, False

        # Must look like a heading style
        is_heading_style = (
            line.isupper()
            or line.istitle()
            or line.endswith(":")
            or re.match(r"^[A-Z][a-z]+([ &/][A-Z][a-z]+)*:?$", line)
        )
        if not is_heading_style:
            return None, False

        # Check alias vocabulary
        if cleaned in SECTION_ALIASES:
            return SECTION_ALIASES[cleaned], True

        # Unrecognised but looks like a header.
        # Guardrails: avoid classifying job title lines with dates/digits as headers.
        if any(c.isdigit() for c in line) or "—" in line:
            return None, False

        # Only keep as misc if not in the ignore list
        if cleaned not in ignore_as_misc:
            return line.rstrip(":").strip(), False

        return None, False

    def _score_confidence(self, result: ParsedResume) -> float:
        score = 0.0

        found = set(result.sections.keys())
        if "skills" in found:
            score += 0.20
        if "experience" in found:
            score += 0.20
        if "education" in found:
            score += 0.10

        text_len = len(result.raw_text)
        if text_len > 2000:
            score += 0.20
        elif text_len > 500:
            score += 0.10

        if result.raw_text:
            non_ascii    = sum(1 for c in result.raw_text if ord(c) > 127)
            control_chars = sum(1 for c in result.raw_text if (ord(c) < 32 and c not in "\n\t\r"))
            symbol_ratio = non_ascii / len(result.raw_text)
            control_ratio = control_chars / len(result.raw_text)
            if "�" in result.raw_text:
                result.warnings.append("Possible encoding issue (replacement characters detected)")
            if control_ratio > 0.02:
                result.warnings.append(f"Possible encoding issue (control character ratio {control_ratio:.1%})")
            if symbol_ratio < 0.05:
                score += 0.20
            elif symbol_ratio < 0.15:
                score += 0.10
            else:
                result.warnings.append(
                    f"High symbol ratio ({symbol_ratio:.1%}) — possible encoding issue"
                )

        if result.layout_type == "single":
            score += 0.10

        return round(min(score, 1.0), 3)


# ── Singleton ─────────────────────────────────────────────────────────────────
resume_parser = ResumeParser()