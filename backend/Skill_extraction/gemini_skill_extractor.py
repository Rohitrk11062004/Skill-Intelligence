#!/usr/bin/env python3
"""
Extract and classify skills from JD .docx files using Gemini.
Outputs a single consolidated JSON file.

Pipeline:
  1. DOCX → section-labeled text (responsibilities, skills, etc.)
  2. LLM proposes skills + categories
  3. normalize_and_recategorize() canonicalizes names, drops junk, overrides categories
  4. Dedup by canonical key
"""

import json
import logging
import os
import re
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

from docx import Document

try:
    # New SDK (recommended). See warning in terminal output.
    from google import genai as genai  # type: ignore
    _GENAI_SDK = "google.genai"
except Exception:  # pragma: no cover
    # Legacy SDK fallback (deprecated upstream, but keeps older envs running).
    import google.generativeai as genai  # type: ignore
    _GENAI_SDK = "google.generativeai"

try:
    from tqdm import tqdm
except ImportError:
    def tqdm(iterable, **kwargs):
        return iterable


def _load_env_file(env_path: Path) -> None:
    """Load KEY=VALUE pairs from an .env file into os.environ if missing."""
    if not env_path.exists() or not env_path.is_file():
        return

    try:
        with env_path.open("r", encoding="utf-8") as f:
            for raw_line in f:
                line = raw_line.strip()
                if not line or line.startswith("#"):
                    continue
                if line.startswith("export "):
                    line = line[len("export "):].strip()
                if "=" not in line:
                    continue

                key, value = line.split("=", 1)
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key and key not in os.environ:
                    os.environ[key] = value
    except Exception:
        # Best-effort loading: ignore malformed .env lines/files.
        return


_SCRIPT_DIR_BOOTSTRAP = Path(__file__).resolve().parent
_PROJECT_ROOT_BOOTSTRAP = _SCRIPT_DIR_BOOTSTRAP.parent
_ENV_CANDIDATES = [
    Path.cwd() / ".env",
    _PROJECT_ROOT_BOOTSTRAP / ".env",
    _SCRIPT_DIR_BOOTSTRAP / ".env",
]

for _env_file in _ENV_CANDIDATES:
    _load_env_file(_env_file)

# ── CONFIGURATION ─────────────────────────────────────────────────────────────

MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent
_DEFAULT_JD_ROOT = _PROJECT_ROOT / "Data"
WORKSPACE_ROOT = Path(os.getenv("JD_ROOT_DIR", str(_DEFAULT_JD_ROOT))).expanduser().resolve()
OUTPUT_FILE = Path(
    os.getenv("SKILL_OUTPUT_FILE", str(WORKSPACE_ROOT / "extracted_skills_results.json"))
).expanduser().resolve()
MAX_RETRIES = 3
INITIAL_BACKOFF_SECONDS = 2

# Extraction strictness toggles (reduce downstream cleanup work)
ENABLE_CO_OCCURRENCE_BOOST = os.getenv("ENABLE_CO_OCCURRENCE_BOOST", "").lower() in ("1", "true", "yes")
ENABLE_DEPARTMENT_FILTER = os.getenv("ENABLE_DEPARTMENT_FILTER", "1").lower() in ("1", "true", "yes")
COLLAPSE_GROUPED_ALTERNATIVES = os.getenv("COLLAPSE_GROUPED_ALTERNATIVES", "1").lower() in ("1", "true", "yes")
DROP_INFERRED_CO_OCCURRENCE = os.getenv("DROP_INFERRED_CO_OCCURRENCE", "1").lower() in ("1", "true", "yes")
MIN_CONFIDENCE = float(os.getenv("MIN_CONFIDENCE", "0.6"))

# ── Observability toggles (default off) ──────────────────────────────────────
LOG_GEMINI_RAW = os.getenv("LOG_GEMINI_RAW", "").lower() in ("1", "true", "yes")
SAVE_GEMINI_RAW_JSON = os.getenv("SAVE_GEMINI_RAW_JSON", "").lower() in ("1", "true", "yes")

SKILL_CATEGORIES = [
    "Technical Skills",
    "Domain/Tools/Process",
    "Team Management",
    "People Management Skills",
    "Communication Skills",
    "Behavioral Skills",
]

GEMINI_RESPONSE_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "role": {"type": "string"},
        "skills": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "category": {"type": "string", "enum": SKILL_CATEGORIES},
                    "confidence": {"type": "number"},
                    "context": {"type": "string"},
                    "is_mandatory": {"type": "boolean"},
                    "importance": {"type": "number"},
                    "prerequisites": {"type": "array", "items": {"type": "string"}},
                    "difficulty": {"type": "integer"},
                    "time_to_learn_hours": {"type": "integer"},
                    "examples": {"type": "array", "items": {"type": "string"}},
                },
                "required": [
                    "name",
                    "category",
                    "confidence",
                    "context",
                    "is_mandatory",
                    "importance",
                    "prerequisites",
                    "difficulty",
                    "time_to_learn_hours",
                    "examples",
                ],
            },
        },
    },
    "required": ["role", "skills"],
}

DOMAIN_SPECIFIC_INDICATORS = [
    "medidata", "rave", "salesforce", "sap", "oracle", "workday",
    "servicenow", "zendesk", "hubspot",
]

ROLE_NAME_NORMALIZATIONS: Dict[str, str] = {
    "sr": "Senior",
    "sde": "Software Development Engineer",
    "sde ii": "Software Development Engineer II",
    "sde iii": "Software Development Engineer III",
    "qa": "Quality Assurance",
    "ml": "Machine Learning",
    "ai": "AI",
    "devops": "DevOps",
}

# ══════════════════════════════════════════════════════════════════════════════
# POST-PROCESSING TABLES
# ══════════════════════════════════════════════════════════════════════════════

# ── Canonical name aliases (key = lowercased input) ───────────────────────────
CANONICAL_NAMES: Dict[str, Optional[str]] = {
    # Languages
    "js": "JavaScript", "javascript": "JavaScript", "typescript": "TypeScript",
    "ts": "TypeScript", "py": "Python", "python3": "Python", "python 3": "Python",
    "c++": "C++", "cpp": "C++", "c#": "C#", "csharp": "C#", "c sharp": "C#",
    "golang": "Go", "go lang": "Go",
    "core java": "Java", "java se": "Java",
    ".net": ".NET", "dot net": ".NET", "dotnet": ".NET",
    # Frontend — canonical form is "React" (not "React.js")
    "reactjs": "React", "react.js": "React", "react js": "React",
    "angularjs": "Angular", "angular.js": "Angular", "angular js": "Angular",
    "vuejs": "Vue.js", "vue": "Vue.js", "vue js": "Vue.js",
    "nextjs": "Next.js", "next js": "Next.js",
    "nuxtjs": "Nuxt.js", "nuxt js": "Nuxt.js",
    # Backend
    "nodejs": "Node.js", "node": "Node.js", "node js": "Node.js",
    "expressjs": "Express.js", "express": "Express.js", "express js": "Express.js",
    "fastapi": "FastAPI", "fast api": "FastAPI",
    "asp.net": "ASP.NET", "aspnet": "ASP.NET",
    "spring boot": "Spring Boot", "springboot": "Spring Boot",
    # Databases
    "postgres": "PostgreSQL", "postgresql": "PostgreSQL",
    "mongo": "MongoDB", "mongodb": "MongoDB", "mongo db": "MongoDB",
    "dynamodb": "DynamoDB", "dynamo db": "DynamoDB",
    "sql server": "SQL Server", "mssql": "SQL Server",
    # Cloud
    "gcp": "Google Cloud", "google cloud platform": "Google Cloud",
    "aws": "AWS", "amazon web services": "AWS",
    "azure": "Microsoft Azure", "microsoft azure": "Microsoft Azure",
    # VCS
    "github": "GitHub", "git hub": "GitHub",
    "gitlab": "GitLab", "git lab": "GitLab",
    "git": "Git",
    "svn": "SVN", "subversion": "SVN",
    "bitbucket": "Bitbucket",
    # CI/CD
    "ci/cd": "CI/CD", "cicd": "CI/CD", "ci cd": "CI/CD", "ci / cd": "CI/CD",
    # Containers
    "k8s": "Kubernetes", "kube": "Kubernetes",
    # AI/ML
    "ml": "Machine Learning", "machine learning": "Machine Learning",
    "dl": "Deep Learning", "deep learning": "Deep Learning",
    "nlp": "NLP", "natural language processing": "NLP",
    "llm": "Large Language Models", "llms": "Large Language Models",
    "gen ai": "Generative AI", "genai": "Generative AI", "generative ai": "Generative AI",
    "gen ai frameworks": "Generative AI Frameworks", "genai frameworks": "Generative AI Frameworks",
    "ocr": "OCR", "optical character recognition": "OCR",
    "computer vision": "Computer Vision", "cv": "Computer Vision",
    "asr": "ASR", "tts": "TTS",
    "conversation ai": "Conversational AI", "conversational ai": "Conversational AI",
    "conversational workflows": "Conversational AI",  # fragment → real skill
    "fine tuning": "Fine-tuning", "fine-tuning": "Fine-tuning", "finetuning": "Fine-tuning",
    "tensorflow": "TensorFlow", "tensor flow": "TensorFlow",
    "pytorch": "PyTorch", "py torch": "PyTorch",
    "scikit-learn": "Scikit-learn", "sklearn": "Scikit-learn",
    "huggingface": "HuggingFace", "hugging face": "HuggingFace",
    "langchain": "LangChain", "lang chain": "LangChain",
    # API integration normalization
    "api integration": "API Integration", "api integrations": "API Integration",
    "api development": "API Development",
    # CS fundamentals
    "oop": "Object-Oriented Programming", "oops": "Object-Oriented Programming",
    "object oriented programming": "Object-Oriented Programming",
    "dsa": "Data Structures & Algorithms",
    "data structures and algorithms": "Data Structures & Algorithms",
    "data structures": "Data Structures & Algorithms",
    # MS Office family
    "ms excel": "MS Excel", "microsoft excel": "MS Excel", "excel": "MS Excel",
    "ms word": "MS Word", "microsoft word": "MS Word", "word": "MS Word",
    "ms powerpoint": "MS PowerPoint", "microsoft powerpoint": "MS PowerPoint",
    "powerpoint": "MS PowerPoint", "ppt": "MS PowerPoint",
    "ms office": "MS Office", "microsoft office suite": "MS Office",
    # Tools
    "vscode": "VS Code", "vs code": "VS Code", "visual studio code": "VS Code",
    "jira": "JIRA", "jenkins": "Jenkins", "postman": "Postman",
    "confluence": "Confluence", "swagger": "Swagger",
    "test link": "TestLink", "testlink": "TestLink",
    # Processes / methodology aliases
    "agile": "Agile", "agile methodology": "Agile", "agile process": "Agile Process",
    "scrum": "Scrum", "kanban": "Kanban",
    "sdlc": "SDLC", "software development life cycle": "SDLC",
    "devops": "DevOps", "dev ops": "DevOps",
    "source code management": "Source Code Management", "scm": "Source Code Management",
    # CS fundamentals aliases
    "multi-threading": "Multi-threading", "multithreading": "Multi-threading",
    "memory management": "Memory Management",
    "data models": "Data Models",
    # iOS / mobile
    "objective c": "Objective C", "objective-c": "Objective C", "objc": "Objective C",
    "view controllers": "View Controllers", "viewcontrollers": "View Controllers",
    "offline storage": "Offline Storage",
    # Behavioral aliases  (Problem Solving → Behavioral per user spec §B.2)
    "decision making": "Decision Making", "decision-making": "Decision Making",
    "problem solving": "Problem Solving", "problem-solving": "Problem Solving",
    # People Management aliases
    "developing team members": "Developing Team Members",
    "feedback - giving": "Feedback - Giving", "feedback giving": "Feedback - Giving",
    "feedback - receiving": "Feedback - Receiving", "feedback receiving": "Feedback - Receiving",
    # VCS aliases
    "knowledge on svn": "SVN", "knowledge of svn": "SVN",
    # Base-name cleanup and canonical normalization additions
    "containerization": "Docker",
    "version control": "Git",
    "cloud infrastructure": "Cloud Platforms",
    "rest api development": "REST API",
    "scripting": "Shell Scripting",
    "frontend": "Frontend Development",
    "frontend development": "Frontend Development",
    "backend": "Backend Development",
    "backend development": "Backend Development",
    "database": "Database",
    "database development": "Database Management",
    "database management": "Database Management",
    "data visualization": "Data Visualization",
    "speech recognition": "Speech Recognition",
    "communication": "Communication",
    # OR/slash alternatives from JD phrasing
    "mongodb/postgresql/mysql": "Database",
    "mongo/postgres/mysql": "Database",
    "node.js/python": "Backend Language",
    "nodejs/python": "Backend Language",
    "aws/gcp/azure": "Cloud Platform",
    "aws/google cloud/azure": "Cloud Platform",
    # Grouped-name cleanup aliases
    "database examples": "Database",
    "backend language examples": "Backend Language",
    "cloud platform examples": "Cloud Platform",
    # Vague best-practices entries are dropped by policy
    "mobile best practices": None,
    "industry best practices": None,
    # Generic fillers to suppress via explicit alias-to-drop
    "technical strategies": None,
    "strong knowledge": None,
    "good knowledge": None,
    "basic knowledge": None,
    "advanced knowledge": None,
    "proven experience": None,
    "extensive experience": None,
    "hands-on experience": None,
    "expertise": None,
    "knowledge": None,
    "experience": None,
    "proficiency": None,
    "understanding": None,
    "familiarity": None,
    "skill": None,
    "skills": None,
}

# ── Forced category overrides (canonical_name → forced category) ──────────────
FORCED_CATEGORY: Dict[str, str] = {
    # ═══ Technical Skills ═════════════════════════════════════════════════════
    # Languages
    "Python": "Technical Skills", "Java": "Technical Skills",
    "JavaScript": "Technical Skills", "TypeScript": "Technical Skills",
    "C++": "Technical Skills", "C#": "Technical Skills", "Go": "Technical Skills",
    "Rust": "Technical Skills", "Ruby": "Technical Skills", "PHP": "Technical Skills",
    "Swift": "Technical Skills", "Kotlin": "Technical Skills", "Scala": "Technical Skills",
    "R": "Technical Skills", "MATLAB": "Technical Skills", "SQL": "Technical Skills",
    "Perl": "Technical Skills", "Shell": "Technical Skills", "Bash": "Technical Skills",
    "PowerShell": "Technical Skills", ".NET": "Technical Skills",
    # Frontend frameworks
    "React": "Technical Skills", "Angular": "Technical Skills", "Vue.js": "Technical Skills",
    "Next.js": "Technical Skills", "Nuxt.js": "Technical Skills",
    "HTML": "Technical Skills", "HTML5": "Technical Skills",
    "CSS": "Technical Skills", "CSS3": "Technical Skills",
    "SASS": "Technical Skills", "SCSS": "Technical Skills",
    "Tailwind CSS": "Technical Skills", "Bootstrap": "Technical Skills",
    "jQuery": "Technical Skills", "Redux": "Technical Skills",
    "Flutter": "Technical Skills", "React Native": "Technical Skills",
    "Streamlit": "Technical Skills",
    # Backend frameworks
    "Node.js": "Technical Skills", "Express.js": "Technical Skills",
    "FastAPI": "Technical Skills", "Django": "Technical Skills", "Flask": "Technical Skills",
    "Spring Boot": "Technical Skills", "Spring": "Technical Skills",
    "ASP.NET": "Technical Skills", "Laravel": "Technical Skills",
    "Rails": "Technical Skills", "Ruby on Rails": "Technical Skills",
    # APIs / Protocols
    "REST API": "Technical Skills", "GraphQL": "Technical Skills",
    "gRPC": "Technical Skills", "WebSocket": "Technical Skills",
    "API Integration": "Technical Skills", "API Development": "Technical Skills",
    # Databases
    "PostgreSQL": "Technical Skills", "MySQL": "Technical Skills",
    "SQLite": "Technical Skills", "MongoDB": "Technical Skills",
    "Redis": "Technical Skills", "Cassandra": "Technical Skills",
    "DynamoDB": "Technical Skills", "Elasticsearch": "Technical Skills",
    "Oracle": "Technical Skills", "SQL Server": "Technical Skills",
    "Firebase": "Technical Skills", "Supabase": "Technical Skills",
    "Neo4j": "Technical Skills", "MariaDB": "Technical Skills",
    # Cloud platforms
    "AWS": "Technical Skills", "Google Cloud": "Technical Skills",
    "Microsoft Azure": "Technical Skills", "Heroku": "Technical Skills",
    "Vercel": "Technical Skills", "Netlify": "Technical Skills",
    "DigitalOcean": "Technical Skills", "Cloudflare": "Technical Skills",
    "AWS SageMaker": "Technical Skills", "AWS Lambda": "Technical Skills",
    "AWS EC2": "Technical Skills", "AWS S3": "Technical Skills",
    # Containerization / Orchestration
    "Docker": "Technical Skills", "Kubernetes": "Technical Skills",
    "Terraform": "Technical Skills", "Ansible": "Technical Skills",
    "Helm": "Technical Skills",
    # AI / ML / NLP / OCR — ALWAYS Technical Skills
    "Machine Learning": "Technical Skills", "Deep Learning": "Technical Skills",
    "NLP": "Technical Skills", "Computer Vision": "Technical Skills",
    "OCR": "Technical Skills", "ASR": "Technical Skills", "TTS": "Technical Skills",
    "Speech-to-Text": "Technical Skills", "Text-to-Speech": "Technical Skills",
    "TensorFlow": "Technical Skills", "PyTorch": "Technical Skills",
    "Keras": "Technical Skills", "Scikit-learn": "Technical Skills",
    "Pandas": "Technical Skills", "NumPy": "Technical Skills",
    "Matplotlib": "Technical Skills", "HuggingFace": "Technical Skills",
    "LangChain": "Technical Skills", "OpenAI": "Technical Skills",
    "Generative AI": "Technical Skills", "Generative AI Frameworks": "Technical Skills",
    "Large Language Models": "Technical Skills",
    "Transformer": "Technical Skills", "RAG": "Technical Skills",
    "FAISS": "Technical Skills", "Conversational AI": "Technical Skills",
    "Chatbot": "Technical Skills", "Fine-tuning": "Technical Skills",
    # CS Fundamentals / Core concepts
    "Data Structures & Algorithms": "Technical Skills",
    "Object-Oriented Programming": "Technical Skills",
    "System Design": "Technical Skills", "Design Patterns": "Technical Skills",
    "Multi-threading": "Technical Skills", "Memory Management": "Technical Skills",
    "Data Models": "Technical Skills", "Microservices": "Technical Skills",
    "Event-Driven Architecture": "Technical Skills",
    "Objective C": "Technical Skills", "View Controllers": "Technical Skills",
    "Offline Storage": "Technical Skills",
    # Processes/methodologies that are technical in nature
    "Source Code Management": "Technical Skills",
    "Agile": "Technical Skills", "Agile Process": "Technical Skills",
    "Scrum": "Technical Skills",
    # Security
    "Cybersecurity": "Technical Skills", "Penetration Testing": "Technical Skills",
    "Firewall": "Technical Skills", "SIEM": "Technical Skills",
    "OWASP": "Technical Skills", "Encryption": "Technical Skills",
    "OAuth": "Technical Skills", "JWT": "Technical Skills",
    "SSL/TLS": "Technical Skills", "IAM": "Technical Skills",
    # Infra / Messaging
    "Kafka": "Technical Skills", "RabbitMQ": "Technical Skills",
    "Celery": "Technical Skills", "Airflow": "Technical Skills",
    "Spark": "Technical Skills", "Hadoop": "Technical Skills",
    "Linux": "Technical Skills", "Unix": "Technical Skills",
    "Nginx": "Technical Skills", "Apache": "Technical Skills",
    "Prometheus": "Technical Skills", "Grafana": "Technical Skills",

    # ═══ Domain/Tools/Process ═════════════════════════════════════════════════
    # VCS / CI-CD / DevOps tools  — Git & GitHub BOTH here for consistency
    "Git": "Domain/Tools/Process", "GitHub": "Domain/Tools/Process",
    "GitLab": "Domain/Tools/Process", "Bitbucket": "Domain/Tools/Process",
    "SVN": "Domain/Tools/Process", "CI/CD": "Domain/Tools/Process",
    "DevOps": "Domain/Tools/Process",
    "GitHub Actions": "Domain/Tools/Process", "GitLab CI": "Domain/Tools/Process",
    "CircleCI": "Domain/Tools/Process", "ArgoCD": "Domain/Tools/Process",
    "Jenkins": "Domain/Tools/Process",
    # Project / Issue tracking
    "JIRA": "Domain/Tools/Process", "Confluence": "Domain/Tools/Process",
    "TestLink": "Domain/Tools/Process", "ServiceNow": "Domain/Tools/Process",
    "Splunk": "Domain/Tools/Process",
    # API / Dev tools
    "Postman": "Domain/Tools/Process", "Swagger": "Domain/Tools/Process",
    "VS Code": "Domain/Tools/Process", "IntelliJ": "Domain/Tools/Process",
    # MS Office
    "MS Excel": "Domain/Tools/Process", "MS Word": "Domain/Tools/Process",
    "MS PowerPoint": "Domain/Tools/Process", "MS Office": "Domain/Tools/Process",
    # Methodologies / Processes
    "Kanban": "Domain/Tools/Process", "SDLC": "Domain/Tools/Process",
    "Waterfall": "Domain/Tools/Process", "SAFe": "Domain/Tools/Process",
    "ITIL": "Domain/Tools/Process", "Six Sigma": "Domain/Tools/Process",
    "MLOps": "Domain/Tools/Process",
    # Compliance / Regulatory
    "ISO 27001": "Domain/Tools/Process", "SOC 2": "Domain/Tools/Process",
    "HIPAA": "Domain/Tools/Process", "GDPR": "Domain/Tools/Process",
    "DPDP Act": "Domain/Tools/Process", "GxP": "Domain/Tools/Process",
    "21 CFR Part 11": "Domain/Tools/Process", "FDA": "Domain/Tools/Process",
    "PCI DSS": "Domain/Tools/Process", "NIST": "Domain/Tools/Process",
    "Rave EDC": "Domain/Tools/Process", "Medidata Rave": "Domain/Tools/Process",
    "Clinical Trials": "Domain/Tools/Process",

    # ═══ Team Management ══════════════════════════════════════════════════════
    "Delegation": "Team Management",
    "Clarity": "Team Management", "Decision Making": "Team Management",
    "Strategic Planning": "Team Management", "Conflict Resolution": "Team Management",
    "Prioritization": "Team Management", "Resource Allocation": "Team Management",
    "Technical Leadership": "Team Management",

    # ═══ People Management Skills ═════════════════════════════════════════════
    "Mentoring": "People Management Skills", "Coaching": "People Management Skills",
    "Onboarding": "People Management Skills", "Hiring": "People Management Skills",
    "Performance Reviews": "People Management Skills",
    "Talent Development": "People Management Skills",
    "Career Development": "People Management Skills",
    "Succession Planning": "People Management Skills",
    "Feedback - Giving": "People Management Skills",
    "Feedback - Receiving": "People Management Skills",
    "Developing Team Members": "People Management Skills",
    "Commitment": "People Management Skills",

    # ═══ Communication Skills ═════════════════════════════════════════════════
    "Stakeholder Communication": "Communication Skills",
    "Presentation Skills": "Communication Skills",
    "Documentation": "Communication Skills",
    "Active Listening": "Communication Skills",
    "Verbal Communication": "Communication Skills",
    "Written Communication": "Communication Skills",
    "Public Speaking": "Communication Skills",
    "Negotiation": "Communication Skills",
    "Facilitation": "Communication Skills",
    "Cross-functional Coordination": "Communication Skills",

    # ═══ Behavioral Skills ════════════════════════════════════════════════════
    "Problem Solving": "Behavioral Skills",          # §B.2 — FIXED
    "Adaptability": "Behavioral Skills",
    "Accountability": "Behavioral Skills",
    "Integrity": "Behavioral Skills",
    "Growth Mindset": "Behavioral Skills",
    "Ownership": "Behavioral Skills",
    "Collaboration": "Behavioral Skills",
    "Time Management": "Behavioral Skills",
    "Resilience": "Behavioral Skills",
    "Self-starter": "Behavioral Skills",
    "Proactiveness": "Behavioral Skills",
    "Empathy": "Behavioral Skills",
    "Emotional Intelligence": "Behavioral Skills",
    "Creativity": "Behavioral Skills",
    "Curiosity": "Behavioral Skills",
    "Work Ethic": "Behavioral Skills",
    # Additional canonical skills introduced by prompt/base-name normalization
    "Database": "Technical Skills",
    "Backend Language": "Technical Skills",
    "Cloud Platform": "Technical Skills",
    "Cloud Platforms": "Technical Skills",
    "Shell Scripting": "Technical Skills",
    "Backend Development": "Technical Skills",
    "Frontend Development": "Technical Skills",
    "Database Management": "Technical Skills",
    "Data Visualization": "Technical Skills",
    "Speech Recognition": "Technical Skills",
    "Analytical Thinking": "Behavioral Skills",
    "Organization": "Behavioral Skills",
    "Communication": "Communication Skills",
}

# ── Keyword-based category rules (regex fallback) ────────────────────────────
KEYWORD_RULES: List[Tuple[str, str]] = [
    # Technical Skills
    (r"\b(api|sdk|framework|library|protocol|database|db|server|runtime)\b", "Technical Skills"),
    (r"\b(programming|language|scripting)\b", "Technical Skills"),
    # Domain / Tools / Process
    (r"\b(compliance|regulation|certification|standard|methodology|audit)\b", "Domain/Tools/Process"),
    (r"\b(iso|soc|hipaa|gdpr|gxp|fda|nist|pci|dpdp)\b", "Domain/Tools/Process"),
    # Team Management
    (r"\b(delegation|prioriti[sz]ation|decision.?making|strategic.?planning|conflict.?resolution|resource.?allocation|sprint.?planning|roadmap|technical.?leadership|clarity)\b", "Team Management"),
    # People Management
    (r"\b(mentor|coach|develop.?team|performance.?review|feedback|talent.?development|succession|onboarding|career.?development|1.on.1|one.on.one|hiring|commitment)\b", "People Management Skills"),
    # Communication Skills
    (r"\b(listening|verbal|written|presentation|public.?speaking|stakeholder.?communication|cross.?functional|negotiation|facilitation|documentation)\b", "Communication Skills"),
    # Behavioral Skills  (problem.?solv moved here per §B.2)
    (r"\b(adaptab|accountab|integrit|growth.?mindset|ownership|collaborat|time.?management|resilience|proactiv|self.?start|self.?motivat|empathy|emotional.?intelligence|creativ|curiosity|work.?ethic|ambiguit|problem.?solv)\b", "Behavioral Skills"),
]

# ── Junk / generic filler to DROP entirely ────────────────────────────────────
# §A.3 Explicit negative list from requirements
DROP_LIST = {
    "coding", "code", "architecture", "product architecture", "product design",
    "deployment",
    "programming", "development", "engineering",
    "troubleshooting", "debugging", "testing", "maintenance",
    "bug fixes", "bug fixing", "application maintenance", "application deployment",
    "service delivery", "project delivery", "process improvement",
    "high-volume transactions", "end-to-end development", "product development",
    "admin dashboards", "cloud infrastructure", "messaging platforms",
    "chatbot frameworks", "internet", "computer", "email",
    "good communication", "team player", "hard working", "self motivated",
    "detail oriented", "fast learner", "quick learner",
    "technical skills", "soft skills", "domain knowledge",
    "system performance monitoring",
    # §A.3 Explicit negative list (must not appear as skills)
    "bug fixes", "troubleshooting", "team members", "conversational workflows",
    "high-volume transactions", "end-to-end development", "project delivery",
    "team members", "pipelines", "scalable solutions", "solutions",
    "best practices", "industry standards", "cross-functional teams",
    "technical documentation", "system design documentation",
    "data pipelines",  # too generic unless qualified
    # Base-name cleanup noise words
    "extensive", "strong", "good", "basic", "advanced", "proven",
    "expertise", "knowledge", "experience", "proficiency", "understanding", "familiarity",
    "frontend/backend development", "backend/frontend development",
    "backend language", "cloud infrastructure management",
    # Standalone modalities for LLM contexts (dropped later in drop_invalid_skill)
    "voice", "image", "speech", "video",
}

# Verb-led phrases = responsibilities, not skills
DROP_LEADING_VERBS = {
    "implement", "build", "develop", "architect", "design", "deliver",
    "drive", "lead", "set", "guide", "ensure", "monitor", "troubleshoot",
    "maintain", "deploy", "collaborate", "manage", "create", "execute",
    "support", "establish", "configure", "analyze", "define", "coordinate",
    "work", "participate", "conduct", "write", "review", "optimize",
    "oversee", "evaluate", "contribute", "facilitate", "provide",
    "identify", "perform", "prepare", "plan", "own", "understand",
}

BASE_NAME_NOISE_WORDS = {
    "extensive", "strong", "good", "basic", "advanced", "proven",
    "expertise", "knowledge", "experience", "proficiency", "skills",
    "skill", "understanding", "familiarity", "of", "in", "with",
}


# ══════════════════════════════════════════════════════════════════════════════
# §B Post-processing functions: canonicalize_skill_name, force_category, drop_invalid_skill
# ══════════════════════════════════════════════════════════════════════════════

def canonicalize_skill_name(name: str, context: str = "") -> Optional[str]:
    """
    §B.1 Canonicalize skill name with alias mapping and modality normalization.
    Apply CANONICAL_NAMES mapping, then modality-aware mapping if in LLM context.
    """
    # Trim, collapse spaces, normalize initial lookup
    name = name.strip()
    name = re.sub(r'\s+', ' ', name)

    # Remove trailing example/examples noise from grouped skill names so
    # the label stays in `examples` field only, not in `name`.
    name = re.sub(r"\s*\bexamples?\b\s*$", "", name, flags=re.IGNORECASE).strip()

    # Base skill extraction: remove common qualifiers and context words.
    cleaned_tokens = [
        t for t in re.split(r"\s+", name)
        if t.strip(".,:;()[]{}").lower() not in BASE_NAME_NOISE_WORDS
    ]
    if cleaned_tokens:
        name = " ".join(cleaned_tokens)
    
    lookup_key = name.lower().strip()
    explicit_hit = lookup_key in CANONICAL_NAMES
    canonical = CANONICAL_NAMES.get(lookup_key, name)

    if canonical is None:
        return None
    
    # If not in CANONICAL_NAMES, apply title-case if appears to be generic
    if canonical == name and not any(c.isupper() for c in name[1:]):
        canonical = name.title()
    
    # §B.1 Modality mapping: apply only when alias table did not already resolve
    # the token. This prevents double-mapping explicit canonical aliases.
    is_llm_context = _is_llm_context(context)
    
    canonical_lower = canonical.lower()
    if is_llm_context and not explicit_hit:
        if canonical_lower in ["voice", "speech", "voice assistant"]:
            canonical = "Speech LLMs"
        elif canonical_lower in ["image", "vision", "image processing"]:
            canonical = "Vision LLMs"
        elif canonical_lower in ["video", "video understanding"]:
            canonical = "Video LLMs"
        elif canonical_lower in ["text", "text generation", "language model"]:
            canonical = "Text LLMs"
    else:
        # Non-LLM contexts: preserve distinct forms
        if canonical_lower == "computer vision":
            canonical = "Computer Vision"
        elif canonical_lower == "speech processing":
            canonical = "Speech Processing"

    # Final guard: ensure trailing example/examples never leak into output name.
    canonical = re.sub(r"\s*\bexamples?\b\s*$", "", canonical, flags=re.IGNORECASE).strip()
    
    return canonical


def force_category(name: str, context: str, current_category: str) -> str:
    """
    §B.2 Apply deterministic category overrides based on name and context.
    Checks FORCED_CATEGORY and KEYWORD_RULES.
    """
    canonical_name = name
    name_lower = name.lower()
    
    # Check FORCED_CATEGORY first (exact + case-insensitive)
    if canonical_name in FORCED_CATEGORY:
        return FORCED_CATEGORY[canonical_name]
    for key, cat in FORCED_CATEGORY.items():
        if key.lower() == name_lower:
            return cat
    
    # Fallback to KEYWORD_RULES (already in force_category via _resolve_category)
    return current_category


def drop_invalid_skill(name: str, context: str) -> bool:
    """
    §B.3 Return True if skill should be dropped, False if valid.
    Check: empty, >6 words, starts with responsibility verbs, banned list, standalone modalities.
    """
    name = name.strip()
    if not name:
        return True
    
    words = name.split()
    if len(words) > 6:
        return True
    
    first_word_lower = words[0].lower()
    if first_word_lower in DROP_LEADING_VERBS:
        return True
    
    name_lower = name.lower()
    if name_lower in DROP_LIST:
        return True

    # Base-skill-only policy: drop vague "best practices" fragments.
    if re.search(r"\bbest practices\b", name_lower):
        return True
    
    # §B.3 Standalone modality tokens for LLM contexts
    is_llm_context = _is_llm_context(context)
    
    if is_llm_context and name_lower in ["voice", "image", "speech", "video"]:
        return True
    
    return False


# ══════════════════════════════════════════════════════════════════════════════
# normalize_and_recategorize (updated to use new functions)
# ══════════════════════════════════════════════════════════════════════════════

def normalize_and_recategorize(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Deterministic post-processing of LLM-proposed skills (§B).
    1. Canonicalize names (§B.1)
    2. Drop invalid skills (§B.3)
    3. Override category (§B.2)
    4. Dedup by canonical key
    """
    seen_keys: Dict[str, int] = {}
    result: List[Dict[str, Any]] = []

    for skill in skills:
        name = str(skill.get("name", "")).strip()
        context = str(skill.get("context", "")).strip()
        if not name:
            continue

        # ── §B.1 Canonicalize name ────────────────────────────────────
        canonical = canonicalize_skill_name(name, context)
        if canonical is None:
            continue

        # ── §B.3 Drop invalid skills ──────────────────────────────────
        if drop_invalid_skill(canonical, context):
            continue

        # ── §B.2 Determine category (forced override → keyword rules → fallback) ──
        proposed = skill.get("category", "Technical Skills")
        final_cat = force_category(canonical, context, proposed)
        if final_cat not in SKILL_CATEGORIES:
            final_cat = _resolve_category(canonical, proposed)
        if final_cat not in SKILL_CATEGORIES:
            final_cat = proposed if proposed in SKILL_CATEGORIES else "Technical Skills"

        # ── Dedup ─────────────────────────────────────────────────────
        dedup_key = canonical.lower().replace(" ", "").replace("-", "").replace(".", "").replace("/", "")
        if dedup_key in seen_keys:
            idx = seen_keys[dedup_key]
            incoming_conf = _confidence_for_compare(skill)
            existing_conf = float(result[idx].get("confidence") or 0.0)
            if incoming_conf > existing_conf:
                result[idx] = _clean(skill, canonical, final_cat)
            continue

        seen_keys[dedup_key] = len(result)
        result.append(_clean(skill, canonical, final_cat))

    result = _apply_llm_specificity_policy(result)
    result = _align_git_github_mandatory(result)

    return result


def _resolve_category(canonical_name: str, proposed: str) -> str:
    """Fallback category resolution using KEYWORD_RULES and FORCED_CATEGORY."""
    name_lower = canonical_name.lower()
    # Try FORCED_CATEGORY first
    for key, cat in FORCED_CATEGORY.items():
        if key.lower() == name_lower:
            return cat
    # Then KEYWORD_RULES
    for pattern, category in KEYWORD_RULES:
        if re.search(pattern, name_lower, re.IGNORECASE):
            return category
    return proposed


def _normalize_context_for_grouping(context: str) -> str:
    c = (context or "").lower().strip()
    c = re.sub(r"\s+", " ", c)
    return c


def _apply_llm_specificity_policy(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Prefer specific LLM modality skills over umbrella LLM in the same context.
    Policy: drop "Large Language Models" when any of Text/Vision/Speech/Video LLMs
    exists for the same normalized context.
    """
    specific_modalities = {"Text LLMs", "Vision LLMs", "Speech LLMs", "Video LLMs"}
    umbrella = "Large Language Models"

    contexts_with_specific = set()
    for s in skills:
        if s.get("name") in specific_modalities:
            contexts_with_specific.add(_normalize_context_for_grouping(str(s.get("context") or "")))

    filtered: List[Dict[str, Any]] = []
    for s in skills:
        if s.get("name") == umbrella:
            key = _normalize_context_for_grouping(str(s.get("context") or ""))
            if key in contexts_with_specific:
                continue
        filtered.append(s)

    return filtered


def _align_git_github_mandatory(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    If Git/GitHub appear in the same normalized context, align is_mandatory and
    importance to the strictest value across those paired rows.
    """
    paired_names = {"Git", "GitHub"}
    grouped: Dict[str, List[int]] = {}

    for idx, s in enumerate(skills):
        if s.get("name") in paired_names:
            key = _normalize_context_for_grouping(str(s.get("context") or ""))
            grouped.setdefault(key, []).append(idx)

    for _, indices in grouped.items():
        if len(indices) < 2:
            continue

        names_in_group = {skills[i].get("name") for i in indices}
        if not {"Git", "GitHub"}.issubset(names_in_group):
            continue

        mandatory_max = any(bool(skills[i].get("is_mandatory")) for i in indices)
        importance_max = max(float(skills[i].get("importance") or 0.0) for i in indices)

        for i in indices:
            skills[i]["is_mandatory"] = mandatory_max
            skills[i]["importance"] = _clamp(importance_max, 0.0, 1.0)

    return skills


def fuzzy_dedup(skills: List[Dict[str, Any]], threshold: float = 0.85) -> List[Dict[str, Any]]:
    """
    Merge skills that are highly similar (Opt #4: Fuzzy Deduplication).
    E.g., 'REST API' vs 'REST APIs', 'React' vs 'ReactJS'.
    Keeps specific tech names (e.g., Python, React) not generics.
    """
    from difflib import SequenceMatcher
    
    result = []
    seen_names = []
    
    for skill in skills:
        name = skill.get("name", "").lower().strip()
        if not name:
            continue
        
        matched_idx = -1
        for idx, existing_name in enumerate(seen_names):
            similarity = SequenceMatcher(None, name, existing_name).ratio()
            if similarity >= threshold:
                matched_idx = idx
                break
        
        if matched_idx >= 0:
            if skill.get("confidence", 0) > result[matched_idx].get("confidence", 0):
                result[matched_idx] = skill
        else:
            result.append(skill)
            seen_names.append(name)
    
    return result


def boost_related_skills(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    If co-occurring skills detected, add related tech (Opt #5: Co-occurrence Boost).
    Keeps specific tech names, not generic grouping.
    E.g., React + JavaScript present -> add HTML, CSS.
    """
    skill_names = {s.get("name", "").lower() for s in skills}
    
    patterns = {
        frozenset(["react", "javascript"]): ["HTML", "CSS"],
        frozenset(["docker", "kubernetes"]): ["Linux"],
        frozenset(["aws", "terraform"]): ["Infrastructure as Code"],
        frozenset(["python", "machine learning"]): ["Data Science"],
        frozenset(["git", "github"]): ["CI/CD"],
    }
    
    added = []
    for pattern, suggested_skills in patterns.items():
        if pattern.issubset(skill_names):
            for suggested in suggested_skills:
                if suggested.lower() not in skill_names and not any(s["name"].lower() == suggested.lower() for s in added):
                    added.append({
                        "name": suggested,
                        "category": "Technical Skills",
                        "confidence": 0.65,
                        "context": "inferred from co-occurrence",
                        "is_mandatory": False,
                        "importance": 0.7,
                        "prerequisites": [],
                        "difficulty": 3,
                        "time_to_learn_hours": 120,
                        "examples": [],
                        "is_domain_specific": False,
                        "domain": "",
                    })
    
    return skills + added


def validate_category_constraints(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Apply category-specific validation (Opt #6: Category Validation).
    Prevents Docker in Behavioral, ensures Git in Domain/Tools/Process, etc.
    """
    generic_banned = {"coding", "programming", "development", "architecture", "deployment"}
    tech_tools = {"python", "javascript", "react", "docker", "jira", "git", "aws", "kubernetes", "sql", "java"}
    
    result = []
    for skill in skills:
        name_lower = skill.get("name", "").lower()
        category = skill.get("category", "")
        
        if category == "Technical Skills" and name_lower in generic_banned:
            continue
        
        if category != "Technical Skills" and any(t in name_lower for t in tech_tools):
            skill["category"] = "Technical Skills"
        
        if name_lower in ["git", "github"] and category != "Domain/Tools/Process":
            skill["category"] = "Domain/Tools/Process"
        
        result.append(skill)
    
    return result


def filter_by_department(skills: List[Dict[str, Any]], department: str) -> List[Dict[str, Any]]:
    """
    Filter irrelevant skills by department (Opt #7: Department-specific Filtering).
    QA: remove ML/AI. Data: remove UI/frontend.
    """
    dept_lower = (department or "").lower()
    
    if "qa" in dept_lower:
        banned = {"machine learning", "deep learning", "nlp", "computer vision", "html", "css", "ui design"}
        return [s for s in skills if s.get("name", "").lower() not in banned]
    
    elif "data" in dept_lower or "analytics" in dept_lower:
        banned = {"html", "css", "ui design", "ux design", "react", "angular", "vue"}
        return [s for s in skills if s.get("name", "").lower() not in banned]
    
    return skills


# ── Operational granularity mapping for role simplification ────────────────────
# When a role has >20 Technical Skills, collapse low-level operational items into
# competency-level groups. These mappings do NOT affect core tech (languages, DBs, clouds, frameworks).
GRANULAR_TO_COMPETENCY = {
    # Access/permissions
    "access permissions": "Access Control",
    "user management": "Access Control",
    "login issues": "Access Control",
    "account management": "Access Control",
    "identity management": "Access Control",
    # Endpoint/antivirus
    "antivirus": "Endpoint Security",
    "endpoint protection": "Endpoint Security",
    "malware protection": "Endpoint Security",
    # Patches
    "security patches": "Patch Management",
    "patch management": "Patch Management",
    "software updates": "Patch Management",
    # Logs/alerts
    "system logs": "Security Monitoring",
    "security alerts": "Security Monitoring",
    "event monitoring": "Security Monitoring",
    "log analysis": "Security Monitoring",
    # Backup/recovery
    "data backups": "Backup & Disaster Recovery",
    "data recovery": "Backup & Disaster Recovery",
    "backup servers": "Backup & Disaster Recovery",
    "disaster recovery": "Backup & Disaster Recovery",
    "backup solutions": "Backup & Disaster Recovery",
    # Troubleshooting
    "hardware troubleshooting": "Troubleshooting & Incident Response",
    "software troubleshooting": "Troubleshooting & Incident Response",
    "system errors": "Troubleshooting & Incident Response",
    "network troubleshooting": "Troubleshooting & Incident Response",
    "incident response": "Troubleshooting & Incident Response",
    # Email
    "email accounts": "Email Management",
    "email issues": "Email Management",
    "email administration": "Email Management",
    "exchange server": "Email Management",
    # Server admin
    "server management": "Server Administration",
    "windows server": "Server Administration",
    "linux administration": "Server Administration",
    "system administration": "Server Administration",
    # Network
    "network configuration": "Network Administration",
    "network maintenance": "Network Administration",
    "firewall": "Network Administration",
    "routers": "Network Administration",
    # Monitoring/Performance
    "system monitoring": "Infrastructure Monitoring",
    "performance monitoring": "Infrastructure Monitoring",
    "resource monitoring": "Infrastructure Monitoring",
    "performance tuning": "Infrastructure Monitoring",
    # Storage
    "storage administration": "Storage Management",
    "storage systems": "Storage Management",
    "san administration": "Storage Management",
    # General IT
    "hardware maintenance": "Technical Support",
    "software installation": "Technical Support",
    "user support": "Technical Support",
    "helpdesk": "Technical Support",
}

# Core skills that NEVER get grouped or dropped
CORE_TECH_WHITELIST = {
    # Specific cloud platforms (NOT generic "cloud platform")
    "aws", "google cloud", "microsoft azure", "gcp",
    # Languages
    "python", "java", "javascript", "typescript", "node.js", "c#", "c++", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "perl", "shell", "bash",
    # Specific databases (NOT generic "database")
    "postgresql", "mysql", "mongodb", "sql server", "oracle", "sqlite", "redis",
    "cassandra", "dynamodb", "elasticsearch", "firebase", "supabase", "mariadb",
    "neo4j", "couchdb", "cockroachdb",
    # Frameworks/libraries (specific, not generic)
    "react", "angular", "vue.js", "next.js", "nuxt.js", "flutter", "react native",
    "fastapi", "django", "flask", "spring boot", "express.js", "rails",
    "aspnet", "laravel", "spring", "hibernate",
    # APIs
    "rest api", "graphql", "grpc", "websocket", "api integration", "api development",
    # Containers/Orchestration
    "docker", "kubernetes", "helm", "openstack",
    # CI/CD
    "jenkins", "gitlab ci", "github actions", "circleci", "argocd",
    # Message queues / Streaming
    "kafka", "rabbitmq", "redis", "activemq",
    # Monitoring/Log aggregation (main platforms)
    "prometheus", "grafana", "elk stack", "datadog", "newrelic",
    # VCS (specific)
    "git", "github", "gitlab", "bitbucket",
    # Search/Analytics
    "elasticsearch", "solr", "splunk",
}

GROUPED_ALTERNATIVE_SKILLS: Dict[str, set[str]] = {
    "Database": {"MongoDB", "PostgreSQL", "MySQL", "SQL Server", "Oracle", "SQLite"},
    "Backend Language": {"Node.js", "Python", "Java", "C#", "Go"},
    "Cloud Platform": {"AWS", "Google Cloud", "Microsoft Azure"},
    "Cloud Platforms": {"AWS", "Google Cloud", "Microsoft Azure"},
}



def group_overly_granular_skills(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Reduce overly-granular IT/admin skill lists by grouping low-level operational
    items into competency-level skills, but only for roles with >20 Technical Skills.
    Core technologies (cloud platforms, languages, databases, frameworks, APIs) are
    never grouped or dropped.
    """
    technical_count = sum(1 for s in skills if s.get("category") == "Technical Skills")
    if technical_count <= 20:
        return skills

    granular_map: Dict[str, List[int]] = {}
    result = []

    for idx, skill in enumerate(skills):
        lower_name = (skill.get("name") or "").lower()

        # Check if this is a core tech that must be preserved
        if lower_name in CORE_TECH_WHITELIST:
            result.append(skill)
        # Check if this skill maps to a competency group
        elif lower_name in GRANULAR_TO_COMPETENCY:
            group_name = GRANULAR_TO_COMPETENCY[lower_name]
            granular_map.setdefault(group_name, []).append(idx)
        else:
            result.append(skill)

    # Merge mapped groups: only add competency row if >=2 items grouped
    grouped_count = 0
    for group_name, indices in granular_map.items():
        if len(indices) >= 2:
            grouped_items = [skills[i] for i in indices]
            competency_row = _merge_grouped_skills(group_name, grouped_items)
            result.append(competency_row)
            grouped_count += 1
        else:
            for i in indices:
                result.append(skills[i])

    if grouped_count > 0:
        before_count = technical_count
        after_tech = sum(1 for s in result if s.get("category") == "Technical Skills")
        logging.info(
            f"[Grouping] Role with >20 technical skills: {before_count} -> {after_tech} "
            f"({len(granular_map)} groups, {grouped_count} competencies added)"
        )

    return result


def _merge_grouped_skills(group_name: str, grouped_items: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Aggregate metadata from grouped items into a single competency row."""
    confidences = [float(s.get("confidence") or 0.0) for s in grouped_items if s.get("confidence") not in (None, "")]
    confidence = max(confidences) if confidences else 0.75
    
    is_mandatory = any(bool(s.get("is_mandatory")) for s in grouped_items)
    
    importances = [float(s.get("importance") or 0.6) for s in grouped_items if s.get("importance") not in (None, "")]
    importance = max(importances) if importances else 0.6
    
    difficulties = [int(s.get("difficulty") or 3) for s in grouped_items if s.get("difficulty")]
    difficulty = max(difficulties) if difficulties else 3
    
    times = [int(s.get("time_to_learn_hours") or 120) for s in grouped_items if s.get("time_to_learn_hours")]
    time_to_learn_hours = max(times) if times else 120

    # Union prerequisites, examples; dedup
    prerequisites: List[str] = []
    examples: List[str] = []
    contexts: List[str] = []
    domain_specific = False
    domain_value = ""

    for item in grouped_items:
        if item.get("prerequisites"):
            prerequisites.extend(item.get("prerequisites", []))
        if item.get("examples"):
            examples.extend(item.get("examples", []))
        ctx = str(item.get("context") or "").strip()
        if ctx:
            contexts.append(ctx)
        if item.get("is_domain_specific"):
            domain_specific = True
            if not domain_value and item.get("domain"):
                domain_value = item.get("domain")

    prerequisites = list(dict.fromkeys(prerequisites))
    examples = list(dict.fromkeys(examples))
    combined_context = " | ".join(contexts)
    if len(combined_context) > 300:
        combined_context = combined_context[:300] + "..."

    return {
        "name": group_name,
        "category": grouped_items[0].get("category", "Technical Skills"),
        "confidence": _clamp(confidence, 0.0, 1.0),
        "context": combined_context,
        "is_mandatory": is_mandatory,
        "importance": _clamp(importance, 0.0, 1.0),
        "prerequisites": prerequisites,
        "difficulty": int(_clamp(float(difficulty), 1.0, 5.0)),
        "time_to_learn_hours": time_to_learn_hours,
        "examples": examples,
        "is_domain_specific": domain_specific,
        "domain": domain_value,
    }


def _skill_key(name: str) -> str:
    return (name or "").lower().replace(" ", "").replace("-", "").replace(".", "").replace("/", "")


def collapse_grouped_alternative_skills(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    If a grouped skill (e.g., Database, Cloud Platform) exists, absorb
    overlapping specific tools into examples and remove redundant rows.
    """
    by_name = {s.get("name"): s for s in skills if s.get("name")}
    names_to_drop = set()

    for grouped_name, member_set in GROUPED_ALTERNATIVE_SKILLS.items():
        grouped_skill = by_name.get(grouped_name)
        if not grouped_skill:
            continue

        existing_examples = _parse_string_list(grouped_skill.get("examples"))
        member_present = [name for name in member_set if name in by_name]
        if member_present:
            merged_examples = list(dict.fromkeys(existing_examples + member_present))
            grouped_skill["examples"] = merged_examples

        for member_name in member_present:
            member = by_name[member_name]
            grouped_skill["is_mandatory"] = bool(grouped_skill.get("is_mandatory")) or bool(member.get("is_mandatory"))
            grouped_skill["importance"] = _clamp(
                max(_safe_float(grouped_skill.get("importance"), 0.6), _safe_float(member.get("importance"), 0.6)),
                0.0,
                1.0,
            )
            grouped_skill["confidence"] = _clamp(
                max(_safe_float(grouped_skill.get("confidence"), 0.75), _safe_float(member.get("confidence"), 0.75)),
                0.0,
                1.0,
            )
            names_to_drop.add(member_name)

    result = [s for s in skills if s.get("name") not in names_to_drop]
    return result


def enforce_final_quality_gate(skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Final deterministic cleanup pass after all transformations.
    Ensures output is canonical, deduped, confidence-filtered, and category-safe.
    """
    dedup_map: Dict[str, Dict[str, Any]] = {}

    for skill in skills:
        name = str(skill.get("name") or "").strip()
        context = str(skill.get("context") or "").strip()

        canonical = canonicalize_skill_name(name, context)
        if canonical is None or drop_invalid_skill(canonical, context):
            continue

        if DROP_INFERRED_CO_OCCURRENCE and context.lower() == "inferred from co-occurrence":
            continue

        conf = _confidence_for_compare(skill)
        if conf < MIN_CONFIDENCE:
            continue

        category = force_category(canonical, context, str(skill.get("category") or "Technical Skills"))
        if category not in SKILL_CATEGORIES:
            category = _resolve_category(canonical, "Technical Skills")

        cleaned = _clean(skill, canonical, category)
        key = _skill_key(canonical)
        existing = dedup_map.get(key)
        if existing is None or _confidence_for_compare(cleaned) > _confidence_for_compare(existing):
            dedup_map[key] = cleaned

    result = list(dedup_map.values())
    result = _apply_llm_specificity_policy(result)
    result = _align_git_github_mandatory(result)
    result.sort(key=lambda s: (not bool(s.get("is_mandatory")), -_safe_float(s.get("importance"), 0.0), s.get("name", "")))
    return result


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _clamp(value: float, min_value: float, max_value: float) -> float:
    """Generic clamp for any numeric type (float or int)."""
    return max(min_value, min(max_value, value))


def _is_llm_context(context: str) -> bool:
    """Check if context indicates LLM/Generative AI scope."""
    llm_keywords = ["llm", "large language model", "generative ai", "gen ai", 
                   "transformer", "foundation model", "gpt", "chatbot"]
    return any(kw in (context or "").lower() for kw in llm_keywords)


def _confidence_for_compare(skill: Dict[str, Any]) -> float:
    raw_conf = skill.get("confidence")
    if raw_conf in (None, ""):
        # Small floor for missing confidence only; explicit 0.0 is preserved.
        return 0.75
    return _clamp(_safe_float(raw_conf, 0.0), 0.0, 1.0)


def _infer_importance_from_context(context: str) -> float:
    ctx = (context or "").lower()
    if any(k in ctx for k in ["must", "required", "mandatory", "must-have", "must have"]):
        return 1.0
    if any(k in ctx for k in ["strong", "proven", "extensive"]):
        return 0.8
    if any(k in ctx for k in ["nice to have", "nice-to-have"]):
        return 0.4
    if any(k in ctx for k in ["good to have", "preferred", "familiarity"]):
        return 0.6
    return 0.6


def _parse_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        v = value.strip().lower()
        if v in {"true", "yes", "1", "y"}:
            return True
        if v in {"false", "no", "0", "n"}:
            return False
    if isinstance(value, (int, float)):
        return bool(value)
    return default


def _parse_string_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        cleaned = [str(v).strip() for v in value if str(v).strip()]
        return list(dict.fromkeys(cleaned))
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []
        parts = [p.strip() for p in re.split(r"[,;/|]", raw) if p.strip()]
        return list(dict.fromkeys(parts))
    return []


def _infer_prerequisites(skill_name: str) -> List[str]:
    n = skill_name.lower()
    if n == "fastapi":
        return ["Python"]
    if n == "react":
        return ["JavaScript", "HTML", "CSS"]
    if n in {"aws sagemaker", "sagemaker"}:
        return ["AWS"]
    if n in {"tensorflow", "pytorch", "scikit-learn", "machine learning framework"}:
        return ["Python", "Mathematics"]
    if n == "kubernetes":
        return ["Docker", "Linux"]
    if n in {"large language models", "llms", "generative ai"}:
        return ["Python", "Machine Learning"]
    if n == "fine-tuning":
        return ["Python", "Machine Learning", "PyTorch", "TensorFlow"]
    return []


# Examples mapping inlined into _clean() for efficiency


def is_domain_specific(skill_name: str) -> tuple[bool, str]:
    lower_name = (skill_name or "").lower()
    for indicator in DOMAIN_SPECIFIC_INDICATORS:
        # Boundary-safe match prevents false positives (e.g. "WhatsApp" != "sap").
        # Example: "WhatsApp Business API" must NOT map to domain "sap".
        if re.search(rf"\b{re.escape(indicator)}\b", lower_name):
            return True, indicator
    return False, ""


def normalize_role_name(raw_role_name: str) -> str:
    base = re.sub(r"[_\-]+", " ", (raw_role_name or "")).strip()
    base = re.sub(r"\s+", " ", base)
    words = base.split(" ")
    normalized_words: List[str] = []
    for w in words:
        key = w.lower()
        normalized_words.append(ROLE_NAME_NORMALIZATIONS.get(key, w.title()))
    return " ".join(normalized_words)


def _clean(original: Dict[str, Any], canonical_name: str, category: str) -> Dict[str, Any]:
    """Build clean skill entry."""
    context = str(original.get("context") or "").strip()

    confidence = _confidence_for_compare(original)
    raw_importance = original.get("importance")
    importance = _clamp(
        _safe_float(raw_importance, _infer_importance_from_context(context)),
        0.0,
        1.0,
    )

    raw_is_mandatory = original.get("is_mandatory")
    if raw_is_mandatory in (None, ""):
        is_mandatory = importance >= 0.8
    else:
        is_mandatory = _parse_bool(raw_is_mandatory, default=importance >= 0.8)

    prerequisites = _parse_string_list(original.get("prerequisites"))
    if not prerequisites:
        prerequisites = _infer_prerequisites(canonical_name)

    difficulty = int(_clamp(float(_safe_float(original.get("difficulty"), 3)), 1.0, 5.0))

    raw_ttl = original.get("time_to_learn_hours")
    if raw_ttl in (None, ""):
        default_hours = {1: 20, 2: 60, 3: 120, 4: 240, 5: 400}[difficulty]
        time_to_learn_hours = default_hours
    else:
        time_to_learn_hours = int(_clamp(float(_safe_float(raw_ttl, 120)), 1.0, 10000.0))

    example_map = {
        "database": ["MongoDB", "PostgreSQL", "MySQL"],
        "backend language": ["Node.js", "Python"],
        "cloud platform": ["AWS", "Google Cloud", "Microsoft Azure"],
        "cloud platforms": ["AWS", "Google Cloud", "Microsoft Azure"],
    }
    examples = _parse_string_list(original.get("examples"))
    if not examples:
        examples = example_map.get(canonical_name.lower(), [])
    domain_specific, domain = is_domain_specific(canonical_name)

    return {
        "name": canonical_name,
        "category": category,
        "confidence": confidence,
        "context": context,
        "is_mandatory": is_mandatory,
        "importance": importance,
        "prerequisites": prerequisites,
        "difficulty": difficulty,
        "time_to_learn_hours": time_to_learn_hours,
        "examples": examples,
        "is_domain_specific": domain_specific,
        "domain": domain,
    }


# ── §D Quality checks ─────────────────────────────────────────────────────────
def run_quality_checks(results: List[Dict[str, Any]]) -> None:
    """
    §D Quality checks to ensure consistency of critical rules.
    Logs warnings if violations found.
    """
    all_skills = []
    for result in results:
        if "skills" in result:
            all_skills.extend(result["skills"])
    
    # §D.1 Git and GitHub always same category
    git_cats = {s["category"] for s in all_skills if s["name"].lower() in ["git", "github"]}
    if len(git_cats) > 1:
        logging.warning(f"[QC] Git/GitHub have inconsistent categories: {git_cats}")
    
    # §D.2 MS Excel/Word/PowerPoint always Domain/Tools/Process
    ms_office_names = {"MS Excel", "MS Word", "MS PowerPoint", "MS Office"}
    for s in all_skills:
        if s["name"] in ms_office_names and s["category"] != "Domain/Tools/Process":
            logging.warning(f"[QC] {s['name']} should be Domain/Tools/Process, got {s['category']}")
    
    # §D.3 Problem Solving always Behavioral Skills
    ps_skills = [s for s in all_skills if s.get("name", "").lower() == "problem solving"]
    for s in ps_skills:
        if s["category"] != "Behavioral Skills":
            logging.warning(f"[QC] Problem Solving should be Behavioral Skills, got {s['category']}")
    
    # §D.4 No banned generic terms
    banned_in_output = {"voice", "image", "speech", "video"} if any(
        "llm" in str(s.get("context", "")).lower() or 
        "generative ai" in str(s.get("context", "")).lower() 
        for s in all_skills
    ) else set()
    
    for s in all_skills:
        if s["name"].lower() in banned_in_output:
            logging.warning(f"[QC] Standalone modality '{s['name']}' in output should have been dropped")


# ══════════════════════════════════════════════════════════════════════════════
# DOCX section-aware parser
# ══════════════════════════════════════════════════════════════════════════════

# Heading keywords → label tag  (checked case-insensitively on short lines ≤8 words)
_SECTION_PATTERNS: List[Tuple[List[str], str]] = [
    # Responsibilities-family
    (["key responsibilities", "responsibilities", "roles and responsibilities",
      "role and responsibilities", "what you will do", "job responsibilities",
      "duties and responsibilities", "primary responsibilities"],
     "[KEY RESPONSIBILITIES]"),
    # Skills-family
    (["required skills", "skills & qualifications", "skills and qualifications",
      "technical skills", "key skills", "must have", "must-have",
      "required qualifications", "requirements", "prerequisites",
      "desired skills", "preferred skills"],
     "[REQUIRED SKILLS]"),
    # Good to have / nice to have
    (["good to have", "nice to have", "preferred qualifications",
      "bonus skills", "additional skills"],
     "[GOOD TO HAVE]"),
    # Soft skills
    (["soft skills", "behavioral competencies", "interpersonal skills",
      "leadership competencies"],
     "[SOFT SKILLS]"),
    # Technology stack (NEW)
    (["technology stack", "tech stack", "technology", "tools & technologies",
      "tools and technologies"],
     "[TECHNOLOGY STACK]"),
    # Company context (NEW)
    (["about the company", "about company", "company overview", "about us",
      "who we are"],
     "[COMPANY_CONTEXT]"),
    # Location / work arrangement (NEW)
    (["location", "work arrangement", "work location", "reporting to",
      "work environment"],
     "[CONTEXT]"),
]


def clean_jd_text(text: str) -> str:
    """
    Remove boilerplate, normalize whitespace, deduplicate identical lines.
    Returns cleaned JD text suitable for LLM processing.
    """
    BOILERPLATE_PATTERNS = [
        r"equal opportunity", r"affirmative action", r"confidential",
        r"benefits include", r"compensation", r"disclaimer",
        r"about (our )?company", r"our mission", r"visit us",
        r"apply now", r"submit (your )?resume", r"careers?\..*\.com",
        r"eoe|equal employ|non-discriminat",
    ]
    
    lines = text.split("\n")
    filtered = []
    seen = set()
    
    for line in lines:
        line = line.strip()
        # Skip empty lines initially (will recombine)
        if not line:
            continue
        # Skip boilerplate
        if any(re.search(p, line, re.IGNORECASE) for p in BOILERPLATE_PATTERNS):
            continue
        # Skip duplicates
        if line not in seen:
            filtered.append(line)
            seen.add(line)
    
    # Normalize whitespace
    text = "\n".join(filtered)
    text = re.sub(r"\n\n+", "\n\n", text)  # collapse multiple newlines
    text = re.sub(r"[\u2022\u2023•]", "-", text)  # normalize bullets
    
    return text.strip()


def optimize_jd_length(labeled_text: str, max_words: int = 5000) -> str:
    """
    If JD exceeds max_words, prioritize key sections to stay within limit.
    Preserves REQUIRED SKILLS first, then KEY RESPONSIBILITIES.
    """
    word_count = len(labeled_text.split())
    if word_count <= max_words:
        return labeled_text
    
    logging.info(f"JD text is {word_count} words (limit {max_words}). Optimizing...")
    
    sections = labeled_text.split("\n\n")
    priority_map = {
        "[REQUIRED SKILLS]": 1,
        "[KEY RESPONSIBILITIES]": 2,
        "[SOFT SKILLS]": 3,
        "[GOOD TO HAVE]": 4,
        "[TECHNOLOGY STACK]": 5,
        "[CONTEXT]": 6,
        "[COMPANY_CONTEXT]": 99,  # lowest priority
    }
    
    def get_priority(section: str) -> int:
        first_line = section.split("\n")[0] if section else ""
        return priority_map.get(first_line, 99)
    
    sections_with_priority = [(get_priority(s), s) for s in sections]
    sections_with_priority.sort(key=lambda x: x[0])
    
    result = []
    for _, section in sections_with_priority:
        result.append(section)
        if len(" ".join(result).split()) >= max_words * 0.9:
            break
    
    return "\n\n".join(result)


def extract_job_metadata(role_name: str, jd_text: str) -> Dict[str, str]:
    """
    Extract seniority level, department, location from role name and JD.
    Returns dict with keys: seniority, department, location.
    """
    metadata = {}
    
    # Seniority detection from role name
    role_lower = role_name.lower()
    if any(x in role_lower for x in ["senior", "sr", "lead", "principal", "staff", "architect"]):
        metadata["seniority"] = "Senior"
    elif any(x in role_lower for x in ["junior", "jr", "graduate", "internship", "intern"]):
        metadata["seniority"] = "Junior"
    else:
        metadata["seniority"] = "Mid-level"
    
    # Department detection from role name
    if any(x in role_lower for x in ["backend", "frontend", "fullstack", "full-stack", "devops", "cloud", "infra"]):
        metadata["department"] = "Engineering"
    elif any(x in role_lower for x in ["qa", "test", "quality"]):
        metadata["department"] = "QA"
    elif any(x in role_lower for x in ["manager", "lead", "director"]):
        metadata["department"] = "Management"
    elif any(x in role_lower for x in ["data", "analyst", "analytics", "bi"]):
        metadata["department"] = "Data/Analytics"
    elif any(x in role_lower for x in ["product"]):
        metadata["department"] = "Product"
    else:
        metadata["department"] = "Engineering"  # default
    
    # Location detection from JD text
    location_patterns = [
        (r"(?:remote|work from home|work from anywhere)", "Remote"),
        (r"(?:on-?site|office)", "On-site"),
        (r"(?:hybrid|flexible)", "Hybrid"),
    ]
    for pattern, label in location_patterns:
        if re.search(pattern, jd_text, re.IGNORECASE):
            metadata["location"] = label
            break
    if "location" not in metadata:
        metadata["location"] = "Not specified"
    
    return metadata


def _detect_section_label(text: str) -> Optional[str]:
    """Return a section label if `text` looks like a heading, else None."""
    words = text.split()
    if len(words) > 8:
        return None
    text_lower = text.lower().rstrip(":").strip()
    for keywords, label in _SECTION_PATTERNS:
        for kw in keywords:
            if kw in text_lower:
                return label
    return None


def read_docx_text(docx_path: Path) -> str:
    """Read .docx, detect sections, clean, optimize length. Return labeled text."""
    try:
        doc = Document(str(docx_path))
        lines: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())

        out_blocks: List[str] = []
        current_block: List[str] = []

        for line in lines:
            text = line.strip()
            label = _detect_section_label(text)

            if label:
                if current_block:
                    out_blocks.append("\n".join(current_block))
                current_block = [label]
            else:
                current_block.append(text)

        if current_block:
            out_blocks.append("\n".join(current_block))

        raw_text = "\n\n".join(out_blocks)
        
        # NEW: Clean boilerplate and duplicates
        raw_text = clean_jd_text(raw_text)
        
        # NEW: Optimize length if needed (prioritize key sections)
        raw_text = optimize_jd_length(raw_text)
        
        return raw_text
    except Exception as exc:
        logging.error(f"Could not read .docx file '{docx_path.name}': {exc}")
        return ""


# ══════════════════════════════════════════════════════════════════════════════
# LLM prompt
# ══════════════════════════════════════════════════════════════════════════════

def build_prompt(role_name: str, jd_text: str) -> str:
    few_shot = """═══ EXTRACTION EXAMPLES (follow this exact style) ═══

Example 1:
JD: "Develop RESTful APIs using Python and Django. MySQL required."
Extracted: Python, Django, REST API, MySQL (all Technical Skills, required)
NOT: "Develop RESTful APIs" (verb), "Experience" (generic)

Example 2:
JD: "AWS, GCP or Azure experience. Strong communication skills."
Extracted: AWS, Google Cloud, Microsoft Azure (Technical Skills), Communication Skills
NOT: "strong" (adjective), "cloud platforms" (generic)

Example 3:
JD: "Kubernetes, Docker, Linux. 5+ years DevOps."
Extracted: Kubernetes, Docker, Linux (Technical Skills, required)
NOT: "deployment" (generic), "5+ years" (not a skill)

Example 4:
JD: "JIRA, Git, Agile, Scrum methodologies."
Extracted: JIRA, Git (Domain/Tools/Process), Agile, Scrum (Technical Skills)
"""
    return f"""
You are an expert HR analyst. Analyze the Job Description for the role: "{role_name}"
Extract ALL relevant skills and classify into EXACTLY these 6 categories.

The JD text below may contain section labels like [KEY RESPONSIBILITIES], [REQUIRED SKILLS], [GOOD TO HAVE], [SOFT SKILLS].
CRITICAL: You MUST extract skills from ALL sections including [KEY RESPONSIBILITIES]. Responsibility bullets often mention specific tools, frameworks, and platforms — capture those as skills.

{few_shot}
═══ EXTRACTION RULES WITH EXAMPLES ═══

═══ WHAT IS A SKILL (noun phrase only) ═══
A skill is a short noun phrase (1–5 words) naming a tool, technology, framework, platform, API, protocol, database, cloud service, AI/ML technique, named regulation/standard, methodology, or personal competency.

CRITICAL: output BASE SKILL NAME ONLY.
- Remove adjectives: extensive, strong, good, basic, advanced, proven, etc.
- Remove context words from the name: expertise, knowledge, experience, proficiency, skills, understanding, familiarity.
- Strip verb-led responsibility text and keep only the noun skill entity.
- Keep non-technical skills too (Team Management, People Management Skills, Communication Skills, Behavioral Skills) with the same base-name cleanup.

GOOD examples: "Python", "REST API", "Docker", "Machine Learning", "MS Excel", "Agile", "Stakeholder Communication", "Problem Solving"
BAD (DO NOT output):
  ✗ "Design and develop scalable microservices" → extract "Microservices" only
  ✗ "Implement REST APIs" → extract "REST API" only
  ✗ "Build CI/CD pipelines" → extract "CI/CD" only
  ✗ "Team Members" → drop (not a skill)
  ✗ "Conversational Workflows" → drop or use "Conversational AI" if context supports it
  ✗ "Pipelines" → drop (too vague; only keep "ML Pipelines" / "Data Pipelines" if explicitly named)
  ✗ "Coding", "Architecture", "Deployment" → drop (generic filler)
  ✗ "Solutions", "Best Practices", "Industry Standards" → drop

═══ BASE NAME NORMALIZATION RULES ═══
- containerization -> Docker
- version control -> Git
- cloud infrastructure -> Cloud Platforms
- API development / REST API development -> REST API
- scripting -> specific language if explicitly stated, else Shell Scripting
- frontend/backend development phrases -> Frontend Development / Backend Development

Bracket list rule:
- "OCR tools (Google Vision API, Tesseract)" -> include OCR, Google Vision API, Tesseract

OR/slash alternatives rule:
- Convert alternatives to ONE domain skill with examples list.
- Example: Database with examples [MongoDB, PostgreSQL, MySQL]
- Example: Backend Language with examples [Node.js, Python]
- Example: Cloud Platform with examples [AWS, Google Cloud, Microsoft Azure]
- Never output skill names containing "Best Practices"; split into concrete skills or omit.

═══ §A.3 EXPLICIT NEGATIVE LIST (must not appear as skills) ═══
MUST DROP: Coding, Architecture, Deployment, Product Architecture, Product Design, Service Delivery, Application Maintenance, Bug Fixes, Troubleshooting, Team Members, Conversational Workflows, High-Volume Transactions, End-to-End Development, Project Delivery

═══ CATEGORY DEFINITIONS ═══

1. "Technical Skills"
   Programming languages, frameworks, libraries, databases, cloud platforms, APIs/protocols, AI/ML/NLP/OCR techniques, containerization, CS fundamentals, Agile/Scrum.
   EXAMPLES: Python, React, Node.js, PostgreSQL, AWS, REST API, OCR, NLP, LLMs, Fine-tuning, Docker, Kubernetes, Microservices, Machine Learning, Conversational AI.
   NOT HERE: MS Excel/Word/PowerPoint (→ Domain/Tools/Process). NOT generic words.

2. "Domain/Tools/Process"
   Software tools, project-tracking platforms, IDEs, office suites, VCS, CI-CD tools, compliance/regulatory standards, process methodologies.
   EXAMPLES: JIRA, Jenkins, Git, GitHub, Confluence, Postman, Swagger, MS Excel, MS Word, MS PowerPoint, MS Office, CI/CD, SDLC, ISO 27001, HIPAA, GDPR.
   Git and GitHub MUST both go here.

3. "Team Management"
   Leading/directing teams, decision-making, strategic thinking, technical leadership.
   EXAMPLES: Delegation, Clarity, Decision Making, Strategic Planning, Conflict Resolution, Prioritization, Resource Allocation, Technical Leadership.
   NOT: mentoring/coaching (→ People Management). NOT: Problem Solving (→ Behavioral).

4. "People Management Skills"
   Developing individuals: mentoring, coaching, feedback, performance management, hiring, onboarding.
   EXAMPLES: Mentoring, Coaching, Hiring, Onboarding, Performance Reviews, Career Development, Feedback - Giving, Feedback - Receiving, Developing Team Members, Commitment.

5. "Communication Skills"
   Specific interpersonal communication competencies.
    EXAMPLES: Active Listening, Verbal Communication, Written Communication, Presentation Skills, Stakeholder Communication, Public Speaking, Documentation, Cross-functional Coordination, Communication.

6. "Behavioral Skills"
   Personal traits, attitudes, character competencies, problem-solving aptitude.
   EXAMPLES: Problem Solving, Adaptability, Accountability, Integrity, Growth Mindset, Ownership, Resilience, Self-starter, Emotional Intelligence, Curiosity, Work Ethic.
   NOT: "Team player" or "Hard working" — too generic, OMIT.

═══ §A.5 MODALITY NORMALIZATION (Generative AI contexts) ═══
If the JD mentions LLM/Large Language Model/Generative AI/Transformer/Foundation Model/GPT, apply modality mapping:
  - "voice" or "speech" → "Speech LLMs"
  - "image" or "vision" → "Vision LLMs"
  - "video" → "Video LLMs"
  - "text generation" or "text" in LLM context → "Text LLMs"
Do NOT output standalone: Voice, Image, Speech, Video (convert or drop).

═══ RULES ═══
1. Extract skills from ALL sections (responsibilities, skills, good-to-have, soft skills).
2. Use ONLY the 6 categories above. Never invent new ones.
3. Normalize: "JS"→"JavaScript", "Postgres"→"PostgreSQL", "ML"→"Machine Learning", "GCP"→"Google Cloud", "React JS"/"React.js"/"ReactJS"→"React", "Conversation AI"→"Conversational AI", "API integrations"→"API Integration", "LLM"→"Large Language Models".
4. JIRA, Jenkins, SVN, Git, GitHub, Confluence, Postman, Swagger, MS Excel, MS Word, MS PowerPoint, CI/CD → always "Domain/Tools/Process".
5. OCR, NLP, LLMs, Fine-tuning, Computer Vision, ASR, TTS, Conversational AI, Speech/Vision/Text/Video LLMs → always "Technical Skills".
6. Problem Solving → always "Behavioral Skills".
7. Mentoring, Coaching, Onboarding, Hiring → always "People Management Skills".
8. Stakeholder Communication, Presentation Skills, Documentation → always "Communication Skills".
9. NEVER output verb-led responsibility phrases. Strip to the entity noun phrase or drop entirely.
10. NEVER output fragments: "Team Members", "Conversational Workflows", "Pipelines", "Solutions".
11. Each skill = short noun phrase (1–5 words).
12. Deduplicate.
13. Return ONLY valid JSON. No markdown fences, no explanation.

═══ ENRICHMENT FIELDS (required for each skill) ═══
- is_mandatory (boolean):
    - true for required/must-have/mandatory
    - false for nice-to-have/familiarity/preferred
- importance (float 0.0-1.0):
    - 1.0 must-have/required
    - 0.8 strong/proven
    - 0.6 good-to-have/familiarity/preferred
    - 0.4 nice-to-have
- prerequisites (array of strings), use these defaults when applicable:
    - FastAPI -> ["Python"]
    - React -> ["JavaScript", "HTML", "CSS"]
    - AWS SageMaker -> ["AWS"]
    - ML framework (TensorFlow/PyTorch/Scikit-learn) -> ["Python", "Mathematics"]
    - Kubernetes -> ["Docker", "Linux"]
    - Large Language Models -> ["Python", "Machine Learning"]
    - Fine-tuning -> ["Python", "Machine Learning", "PyTorch", "TensorFlow"]
    - Otherwise []
- difficulty (integer 1-5):
    - 1 very easy, 2 easy, 3 moderate, 4 hard, 5 expert-level
- time_to_learn_hours (integer): rough estimate aligned to difficulty
- examples (array):
    - usually populated for OR/slash grouped domain skills
    - usually [] for simple non-grouped skills and most non-technical skills

═══ OUTPUT FORMAT ═══
{{
  "role": "{role_name}",
  "skills": [
    {{
      "name": "<canonical skill name>",
      "category": "<one of the 6 exact categories>",
            "confidence": 0.85,
            "context": "<exact phrase from JD>",
            "is_mandatory": true,
            "importance": 1.0,
            "prerequisites": ["<skill>", "<skill>"],
            "difficulty": 3,
            "time_to_learn_hours": 120,
            "examples": ["<optional example>"]
    }}
  ]
}}

Return STRICT JSON only. Do not include markdown fences, comments, or extra keys.

═══ JOB DESCRIPTION ═══
{jd_text}
"""


def build_compact_retry_prompt(role_name: str, jd_text: str) -> str:
        """Short fallback prompt for stubborn JDs that fail JSON formatting."""
        return f"""
Return STRICT JSON only, no markdown, no comments.

Extract skills from this JD for role "{role_name}".
Use exactly these categories:
- Technical Skills
- Domain/Tools/Process
- Team Management
- People Management Skills
- Communication Skills
- Behavioral Skills

Rules:
1. Skill names must be short noun phrases (1-5 words).
2. Deduplicate skills.
3. Use canonical names (e.g., React, Node.js, REST API, Google Cloud, Machine Learning, CI/CD).
4. Do not include generic fillers like "experience", "knowledge", "best practices".
5. Return fields exactly as below.

Output format:
{{
    "role": "{role_name}",
    "skills": [
        {{
            "name": "<skill>",
            "category": "<one of 6 categories>",
            "confidence": 0.85,
            "context": "<jd phrase>",
            "is_mandatory": true,
            "importance": 1.0,
            "prerequisites": [],
            "difficulty": 3,
            "time_to_learn_hours": 120,
            "examples": []
        }}
    ]
}}

JD:
{jd_text}
"""


def build_line_retry_prompt(role_name: str, jd_text: str) -> str:
    """Last-resort fallback prompt that avoids JSON entirely."""
    return f"""
Extract skills from this JD for role "{role_name}".

Return plain text only. One skill per line in this exact format:
name|category|is_mandatory|importance|context

Rules:
1. category must be one of:
   Technical Skills
   Domain/Tools/Process
   Team Management
   People Management Skills
   Communication Skills
   Behavioral Skills
2. is_mandatory must be true or false.
3. importance must be a float between 0.0 and 1.0.
4. Keep context short (<= 80 chars), no pipe character.
5. No markdown, no bullets, no numbering.

JD:
{jd_text}
"""


# ══════════════════════════════════════════════════════════════════════════════
# LLM call + pipeline
# ══════════════════════════════════════════════════════════════════════════════

def strip_markdown(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"^\s*```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```\s*$", "", cleaned)
    return cleaned.strip()


def _extract_first_json_object(text: str) -> str:
    """Extract the first balanced JSON object from model output."""
    start = text.find("{")
    if start == -1:
        return text

    depth = 0
    in_string = False
    escape = False

    for idx in range(start, len(text)):
        ch = text[idx]
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[start:idx + 1]

    return text[start:]


def _repair_json_text(raw_text: str) -> str:
    """Best-effort cleanup for common LLM JSON mistakes."""
    candidate = strip_markdown(raw_text)
    candidate = _extract_first_json_object(candidate)

    # Remove JS-style comments.
    candidate = re.sub(r"/\*.*?\*/", "", candidate, flags=re.DOTALL)
    candidate = re.sub(r"(^|\s)//.*?$", "", candidate, flags=re.MULTILINE)

    # Normalize smart quotes to standard quotes.
    candidate = (
        candidate
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2018", "'")
        .replace("\u2019", "'")
    )

    # Remove trailing commas before object/array close.
    candidate = re.sub(r",\s*([}\]])", r"\1", candidate)
    return candidate.strip()


def _parse_gemini_json(raw_text: str, jd_stem: str = "") -> Dict[str, Any]:
    """Parse Gemini output with fallback cleanup steps."""
    first_pass = strip_markdown(raw_text)
    try:
        return json.loads(first_pass)
    except json.JSONDecodeError:
        repaired = _repair_json_text(raw_text)
        try:
            return json.loads(repaired)
        except json.JSONDecodeError as repair_err:
            if jd_stem:
                logging.warning(f"[ParseFallback] JSON parse failed for {jd_stem}: {repair_err}")
            raise


def _parse_line_fallback_output(raw_text: str, role_name: str) -> Dict[str, Any]:
    """Parse pipe-delimited fallback output into standard payload shape."""
    skills: List[Dict[str, Any]] = []

    for raw_line in raw_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("#"):
            continue

        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 5:
            continue

        name, category, mandatory_raw, importance_raw, context = parts[:5]
        if not name:
            continue

        if category not in SKILL_CATEGORIES:
            category = "Technical Skills"

        is_mandatory = mandatory_raw.lower() in {"true", "1", "yes", "y"}
        importance = _clamp(_safe_float(importance_raw, 0.6), 0.0, 1.0)

        skills.append({
            "name": name,
            "category": category,
            "confidence": 0.75,
            "context": context.replace("|", "/")[:200],
            "is_mandatory": is_mandatory,
            "importance": importance,
            "prerequisites": [],
            "difficulty": 3,
            "time_to_learn_hours": 120,
            "examples": [],
        })

    return {"role": role_name, "skills": skills}


def call_gemini_line_fallback(
    model: Any,
    prompt: str,
    role_name: str,
    jd_stem: str = "",
) -> Optional[Dict[str, Any]]:
    """Last-resort model call using plain text lines instead of JSON."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            raw_text = _generate_text(model=model, prompt=prompt, generation_config={
                "temperature": 0.0,
                "max_output_tokens": 1536,
                "top_p": 0.9,
            })
            if not raw_text.strip():
                continue

            parsed = _parse_line_fallback_output(raw_text, role_name=role_name)
            if parsed.get("skills"):
                return parsed
        except Exception as e:
            if attempt < MAX_RETRIES:
                time.sleep(INITIAL_BACKOFF_SECONDS * (2 ** (attempt - 1)))
            else:
                logging.error(f"Line fallback failed after {MAX_RETRIES} attempts ({jd_stem}): {e}")
    return None


def call_gemini(
    model: Any,
    prompt: str,
    jd_stem: str = "",
    max_output_tokens: int = 4096,
) -> Optional[Dict[str, Any]]:
    """Call Gemini with optimized generation config for consistency (Opt #2)."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            # Temperature control: 0.1 = deterministic/consistent (not creative)
            raw_text = _generate_text(model=model, prompt=prompt, generation_config={
                "temperature": 0.1,      # Low = consistent output
                "max_output_tokens": max_output_tokens,
                "top_p": 0.95,
                "response_mime_type": "application/json",
                "response_schema": GEMINI_RESPONSE_SCHEMA,
            })
            if not raw_text:
                continue

            # ── Observability ─────────────────────────────────────────────
            if LOG_GEMINI_RAW:
                logging.info(f"[RAW] {jd_stem}:\n{raw_text[:2000]}")

            parsed = _parse_gemini_json(raw_text, jd_stem=jd_stem)

            if SAVE_GEMINI_RAW_JSON and jd_stem:
                raw_path = OUTPUT_FILE.parent / f"_raw_{jd_stem}.json"
                with open(raw_path, "w", encoding="utf-8") as rf:
                    json.dump(parsed, rf, indent=2, ensure_ascii=False)
                logging.info(f"[RAW-SAVE] {raw_path.name}")

            return parsed
        except Exception as e:
            if attempt < MAX_RETRIES:
                time.sleep(INITIAL_BACKOFF_SECONDS * (2 ** (attempt - 1)))
            else:
                logging.error(f"Gemini API Error after {MAX_RETRIES} attempts: {e}")
    return None


def _generate_text(model: Any, prompt: str, generation_config: Dict[str, Any]) -> str:
    """
    Thin adapter over both SDKs:
    - google.genai: client.models.generate_content(...).text
    - google.generativeai: GenerativeModel.generate_content(...).text
    """
    if _GENAI_SDK == "google.genai":
        # google.genai client: model is the Client instance
        response = model.models.generate_content(
            model=MODEL_NAME,
            contents=prompt,
            config=generation_config,
        )
        return getattr(response, "text", None) or ""

    # Legacy google.generativeai
    response = model.generate_content(prompt, generation_config=generation_config)
    return getattr(response, "text", None) or ""


def main():
    logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)

    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if not api_key:
        logging.error("Neither GEMINI_API_KEY nor GOOGLE_API_KEY environment variable is set.")
        return 1

    if _GENAI_SDK == "google.genai":
        # New SDK: create a client (model name is passed per-call).
        model = genai.Client(api_key=api_key)
        logging.info(f"Using Gemini SDK: {_GENAI_SDK} (model={MODEL_NAME})")
    else:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(MODEL_NAME)
        logging.info(f"Using Gemini SDK: {_GENAI_SDK} (model={MODEL_NAME})")

    docx_files = [p for p in WORKSPACE_ROOT.rglob("*.docx")
                  if p.is_file()
                  and not p.name.startswith("~$")
                  and p.name.lower() != "default.docx"
                  and ".venv" not in str(p)
                  and "node_modules" not in str(p)]

    if not docx_files:
        logging.warning(f"No .docx files found in {WORKSPACE_ROOT}")
        return 0

    logging.info(f"Found {len(docx_files)} JD files. Starting extraction...")

    final_results = []
    for docx_file in tqdm(docx_files, desc="Extracting Skills"):
        text = read_docx_text(docx_file)
        if not text or len(text) < 50:
            logging.warning(f"Skipping {docx_file.name} - empty or too short.")
            continue

        result = call_gemini(model, build_prompt(docx_file.stem, text),
                             jd_stem=docx_file.stem)
        if not result:
            # Retry with a shorter prompt + shorter JD slice to avoid malformed long outputs.
            compact_text = text[:12000]
            result = call_gemini(
                model,
                build_compact_retry_prompt(docx_file.stem, compact_text),
                jd_stem=f"{docx_file.stem}-retry",
                max_output_tokens=2048,
            )
            if result:
                logging.info(f"[RetrySuccess] Recovered extraction for {docx_file.stem} with compact retry.")

        if not result:
            # Last-resort fallback that avoids JSON output formatting issues.
            compact_text = text[:9000]
            result = call_gemini_line_fallback(
                model,
                build_line_retry_prompt(docx_file.stem, compact_text),
                role_name=docx_file.stem,
                jd_stem=f"{docx_file.stem}-line-retry",
            )
            if result:
                logging.info(f"[LineRetrySuccess] Recovered extraction for {docx_file.stem} with line fallback.")

        if result:
            result["role"] = normalize_role_name(docx_file.stem)
            result["file_path"] = str(docx_file.relative_to(WORKSPACE_ROOT))
            result["folder_category"] = docx_file.parent.name
            
            # NEW: Extract and add job metadata (seniority, department, location)
            job_metadata = extract_job_metadata(docx_file.stem, text)
            result["metadata"] = job_metadata

            if "skills" in result:
                raw_count = len(result["skills"])
                result["skills"] = normalize_and_recategorize(result["skills"])
                clean_count = len(result["skills"])
                
                # NEW: Post-processing optimizations
                result["skills"] = fuzzy_dedup(result["skills"], threshold=0.85)  # Opt #4
                result["skills"] = validate_category_constraints(result["skills"])  # Opt #6
                if ENABLE_CO_OCCURRENCE_BOOST:
                    result["skills"] = boost_related_skills(result["skills"])  # Opt #5
                if ENABLE_DEPARTMENT_FILTER and job_metadata.get("department"):
                    result["skills"] = filter_by_department(result["skills"], job_metadata["department"])  # Opt #7
                
                result["skills"] = group_overly_granular_skills(result["skills"])
                if COLLAPSE_GROUPED_ALTERNATIVES:
                    result["skills"] = collapse_grouped_alternative_skills(result["skills"])
                result["skills"] = enforce_final_quality_gate(result["skills"])
                grouped_count = len(result["skills"])
                result["total_skills_found"] = grouped_count
                if raw_count != clean_count:
                    logging.info(
                        f"  {docx_file.stem}: {raw_count} raw -> {clean_count} clean "
                        f"({raw_count - clean_count} dropped)"
                    )
                if grouped_count != clean_count:
                    logging.info(
                        f"  {docx_file.stem}: grouped operational skills {clean_count} -> {grouped_count}"
                    )

            final_results.append(result)
        else:
            final_results.append({
                "role": docx_file.stem,
                "file_path": str(docx_file.relative_to(WORKSPACE_ROOT)),
                "error": "Failed to extract from this file",
                "skills": []
            })

    # ── §D Quality checks ─────────────────────────────────────────────────────
    run_quality_checks(final_results)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(final_results, f, indent=2, ensure_ascii=False)

    logging.info(f"Successfully processed {len(final_results)} files.")
    logging.info(f"Consolidated results saved to: {OUTPUT_FILE}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
