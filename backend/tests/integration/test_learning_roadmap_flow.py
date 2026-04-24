"""Integration tests for the Week 4 learning roadmap flow."""
from __future__ import annotations

import asyncio
from pathlib import Path
import json
import math
import re
from unittest.mock import AsyncMock
import uuid

import httpx
import pytest
import pytest_asyncio
from docx import Document
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.pool import StaticPool

import app.db.session as db_session
import app.main as main_module
from app.api.v1.endpoints import assessments as assessments_endpoint
from app.api.v1.endpoints import learning as learning_endpoint
from app.api.v1.endpoints import week_assessments as week_assessments_endpoint
from app.models.models import Base, LearningPlanItem, LearningPlanItemResource, Role, RoleSkillRequirement, Skill, WeekAssessment, WeekAssessmentAttempt
from app.schemas.learning import SubSubtopic, Subtopic
from app.services.extraction import llm_extractor, regex_extractor
from app.services.gap.gap_detector import GapItem, GapResult
from app.services.learning import path_generator
from app.services import week_assessment_service

TARGET_ROLE_NAME = "Senior Software Engineer"
RESUME_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


async def _mock_generate_assessment_questions(skill_name: str, level: str, subtopics: list[str]) -> list[dict]:
    _ = skill_name
    _ = level
    _ = subtopics
    return [
        {
            "question": "Question 1",
            "options": ["A", "B", "C", "D"],
            "correct_index": 0,
            "explanation": "Explanation 1",
        },
        {
            "question": "Question 2",
            "options": ["A", "B", "C", "D"],
            "correct_index": 1,
            "explanation": "Explanation 2",
        },
        {
            "question": "Question 3",
            "options": ["A", "B", "C", "D"],
            "correct_index": 2,
            "explanation": "Explanation 3",
        },
        {
            "question": "Question 4",
            "options": ["A", "B", "C", "D"],
            "correct_index": 3,
            "explanation": "Explanation 4",
        },
        {
            "question": "Question 5",
            "options": ["A", "B", "C", "D"],
            "correct_index": 0,
            "explanation": "Explanation 5",
        },
    ]


async def _mock_generate_week_assessment_questions(
    *,
    week_number: int,
    skills: list[str],
    subtopics: list[str],
    question_count: int,
) -> list[dict]:
    _ = week_number
    _ = subtopics
    lead_skill = skills[0] if skills else "General"
    questions: list[dict] = []
    for idx in range(question_count):
        questions.append(
            {
                "question": f"Week question {idx + 1}",
                "options": ["A", "B", "C", "D"],
                "correct_index": idx % 4,
                "explanation": f"Week explanation {idx + 1}",
                "tags": [
                    f"skill:{lead_skill}",
                    "subtopic:Foundations",
                    "type:edge_case" if idx % 5 == 0 else "type:scenario",
                ],
            }
        )
    return questions


async def _mock_generate_skill_breakdown(
    skill_name: str,
    target_role: str,
    gap_type: str,
    time_to_learn_hours: float,
    user_skill_profile: list[tuple[str, str]] | None = None,
) -> list[Subtopic]:
    _ = skill_name
    _ = target_role
    _ = gap_type
    _ = user_skill_profile
    total_hours = float(time_to_learn_hours or 0.0)
    return [
        Subtopic(
            title="Test Fundamentals",
            estimated_hours=total_hours * 0.4,
            sub_subtopics=[
                SubSubtopic(title="Core concepts", estimated_hours=total_hours * 0.2),
                SubSubtopic(title="Practice", estimated_hours=total_hours * 0.2),
            ],
        ),
        Subtopic(
            title="Test Applied",
            estimated_hours=total_hours * 0.3,
            sub_subtopics=[
                SubSubtopic(title="Exercises", estimated_hours=total_hours * 0.15),
                SubSubtopic(title="Projects", estimated_hours=total_hours * 0.15),
            ],
        ),
        Subtopic(
            title="Test Advanced",
            estimated_hours=total_hours * 0.3,
            sub_subtopics=[
                SubSubtopic(title="Deep dive", estimated_hours=total_hours * 0.15),
                SubSubtopic(title="Real world", estimated_hours=total_hours * 0.15),
            ],
        ),
    ]


async def _mock_generate_assessment_breakdown(
    skill_name: str,
    target_role: str,
    failed_areas: list[str],
    time_to_learn_hours: float,
) -> list[Subtopic]:
    _ = skill_name
    _ = target_role
    _ = failed_areas
    total_hours = float(time_to_learn_hours or 0.0)
    return [
        Subtopic(
            title="Remediation Fundamentals",
            estimated_hours=total_hours * 0.34,
            sub_subtopics=[
                SubSubtopic(title="Core concepts", estimated_hours=total_hours * 0.17),
                SubSubtopic(title="Practice", estimated_hours=total_hours * 0.17),
            ],
        ),
        Subtopic(
            title="Remediation Practice",
            estimated_hours=total_hours * 0.33,
            sub_subtopics=[
                SubSubtopic(title="Exercises", estimated_hours=total_hours * 0.165),
                SubSubtopic(title="Drills", estimated_hours=total_hours * 0.165),
            ],
        ),
        Subtopic(
            title="Remediation Application",
            estimated_hours=total_hours * 0.33,
            sub_subtopics=[
                SubSubtopic(title="Projects", estimated_hours=total_hours * 0.165),
                SubSubtopic(title="Review", estimated_hours=total_hours * 0.165),
            ],
        ),
    ]


@pytest_asyncio.fixture
async def async_client(monkeypatch: pytest.MonkeyPatch):
    test_engine = create_async_engine(
        "sqlite+aiosqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    test_session_factory = async_sessionmaker(
        bind=test_engine,
        class_=AsyncSession,
        expire_on_commit=False,
        autoflush=False,
    )

    monkeypatch.setattr(db_session, "engine", test_engine)
    monkeypatch.setattr(db_session, "AsyncSessionLocal", test_session_factory)
    monkeypatch.setattr(main_module, "engine", test_engine)
    monkeypatch.setattr(main_module, "AsyncSessionLocal", test_session_factory)

    monkeypatch.setattr(regex_extractor, "extract_from_sections", lambda sections: [])
    monkeypatch.setattr(llm_extractor, "extract", AsyncMock(return_value=[]))
    monkeypatch.setattr(path_generator, "_generate_skill_breakdown", _mock_generate_skill_breakdown)
    monkeypatch.setattr(path_generator, "_generate_assessment_breakdown", _mock_generate_assessment_breakdown)
    monkeypatch.setattr(learning_endpoint, "generate_roadmap", path_generator.generate_roadmap)
    monkeypatch.setattr(
        week_assessments_endpoint,
        "generate_week_assessment_questions",
        _mock_generate_week_assessment_questions,
    )

    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    async with main_module.app.router.lifespan_context(main_module.app):
        transport = httpx.ASGITransport(app=main_module.app)
        async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
            yield client

    await test_engine.dispose()


async def _create_user_and_token(client: httpx.AsyncClient, suffix: str) -> tuple[str, str, dict[str, str]]:
    email = f"roadmap_{suffix}@example.com"
    password = "StrongPass123!"

    register_response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "password": password,
            "full_name": f"Roadmap User {suffix}",
            "department": "Engineering",
            "job_title": "Software Engineer",
        },
    )
    assert register_response.status_code == 201, register_response.text

    login_response = await client.post(
        "/api/v1/auth/login",
        data={"username": email, "password": password},
        headers={"Content-Type": "application/x-www-form-urlencoded"},
    )
    assert login_response.status_code == 200, login_response.text
    token = login_response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    me_response = await client.get("/api/v1/auth/me", headers=headers)
    assert me_response.status_code == 200, me_response.text
    return me_response.json()["id"], email, headers


async def _seed_target_role(suffix: str) -> str:
    async with db_session.AsyncSessionLocal() as db:
        role = Role(
            name=TARGET_ROLE_NAME,
            description=f"Learning roadmap target role {suffix}",
            department=None,
            seniority_level="Senior",
            is_custom=True,
        )
        db.add(role)
        await db.flush()

        skill_one = Skill(
            name=f"Roadmap Skill A {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=3,
            time_to_learn_hours=60,
            source_role=role.name,
        )
        skill_two = Skill(
            name=f"Roadmap Skill B {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=4,
            time_to_learn_hours=90,
            source_role=role.name,
        )
        db.add_all([skill_one, skill_two])
        await db.flush()

        db.add_all(
            [
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_one.id,
                    importance=1.0,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_two.id,
                    importance=0.8,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
            ]
        )
        await db.commit()
        return role.id


async def _seed_prereq_chain_role(suffix: str) -> str:
    async with db_session.AsyncSessionLocal() as db:
        role = Role(
            name=TARGET_ROLE_NAME,
            description=f"Prereq chain role {suffix}",
            department=None,
            seniority_level="Senior",
            is_custom=True,
        )
        db.add(role)
        await db.flush()

        skill_a = Skill(
            name=f"Chain Skill A {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=2,
            time_to_learn_hours=12,
            source_role=role.name,
        )
        skill_b = Skill(
            name=f"Chain Skill B {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites=json.dumps([f"Chain Skill A {suffix}"]),
            difficulty=3,
            time_to_learn_hours=14,
            source_role=role.name,
        )
        skill_d = Skill(
            name=f"Chain Skill D {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=4,
            time_to_learn_hours=20,
            source_role=role.name,
        )
        db.add_all([skill_a, skill_b, skill_d])
        await db.flush()

        db.add_all(
            [
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_a.id,
                    importance=1.0,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_b.id,
                    importance=1.2,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_d.id,
                    importance=0.5,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
            ]
        )
        await db.commit()
        return role.id


def _create_resume_file(tmp_path: Path, suffix: str) -> Path:
    document = Document()
    document.add_paragraph("John Doe")
    document.add_paragraph("john.doe@example.com")
    document.add_paragraph("")
    document.add_paragraph("PROFESSIONAL SUMMARY")
    document.add_paragraph("Software engineer with backend and cloud experience.")
    document.add_paragraph("")
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("Software Engineer — Example Corp (2021-2024)")
    document.add_paragraph("Built REST APIs with Python and FastAPI.")
    document.add_paragraph("Managed PostgreSQL databases and Docker deployments.")
    document.add_paragraph("")
    document.add_paragraph("TECHNICAL SKILLS")
    document.add_paragraph("Python, FastAPI, PostgreSQL, Docker, AWS, Redis, Git")
    document.add_paragraph("")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("B.Tech Computer Science")

    file_path = tmp_path / f"resume_{suffix}.docx"
    document.save(file_path)
    return file_path


async def _upload_and_process_resume(client: httpx.AsyncClient, headers: dict[str, str], tmp_path: Path, suffix: str) -> str:
    resume_file = _create_resume_file(tmp_path, suffix)
    with resume_file.open("rb") as file_handle:
        upload_response = await client.post(
            "/api/v1/resumes/upload",
            headers=headers,
            files={"file": (resume_file.name, file_handle, RESUME_MIME_TYPE)},
        )
    assert upload_response.status_code == 202, upload_response.text
    job_id = upload_response.json()["job_id"]

    process_response = await client.post(f"/api/v1/resumes/{job_id}/process", headers=headers)
    assert process_response.status_code == 200, process_response.text
    return job_id


async def _set_target_role(client: httpx.AsyncClient, headers: dict[str, str], role_name: str) -> None:
    response = await client.post(
        "/api/v1/users/me/target-role",
        headers=headers,
        json={"role_name": role_name},
    )
    assert response.status_code == 200, response.text


async def _generate_roadmap(
    client: httpx.AsyncClient,
    headers: dict[str, str],
    hours_per_week: float = 10.0,
    deadline_weeks: int | None = None,
    daily_hours: float | None = None,
) -> dict:
    params = [f"hours_per_week={hours_per_week}"]
    if deadline_weeks is not None:
        params.append(f"deadline_weeks={deadline_weeks}")
    if daily_hours is not None:
        params.append(f"daily_hours={daily_hours}")
    response = await client.get(f"/api/v1/users/me/learning-roadmap?{'&'.join(params)}", headers=headers)
    assert response.status_code == 200, response.text
    return response.json()


async def _prepare_learning_flow(client: httpx.AsyncClient, tmp_path: Path, suffix: str) -> tuple[dict, str, str, dict[str, str]]:
    _, _, headers = await _create_user_and_token(client, suffix)
    role_id = await _seed_target_role(suffix)
    await _upload_and_process_resume(client, headers, tmp_path, suffix)
    await _set_target_role(client, headers, TARGET_ROLE_NAME)
    roadmap = await _generate_roadmap(client, headers)
    plan_id = roadmap["plan_id"]
    item_id = roadmap["weeks"][0]["skills"][0]["item_id"]
    return roadmap, plan_id, item_id, headers


async def _complete_first_item(client: httpx.AsyncClient, headers: dict[str, str], plan_id: str, item_id: str) -> tuple[dict, dict]:
    in_progress_response = await client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "in_progress"},
    )
    assert in_progress_response.status_code == 200, in_progress_response.text

    completed_response = await client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "completed"},
    )
    assert completed_response.status_code == 200, completed_response.text
    return in_progress_response.json(), completed_response.json()


async def _clear_learning_state(client: httpx.AsyncClient, headers: dict[str, str]) -> None:
    _ = client
    _ = headers


async def _create_fresh_user(client: httpx.AsyncClient, suffix: str) -> dict[str, str]:
    _, _, headers = await _create_user_and_token(client, suffix)
    return headers


def _assert_roadmap_payload(payload: dict) -> tuple[str, str]:
    assert payload["plan_id"]
    assert payload["target_role"]
    assert "readiness_score" in payload
    assert payload["total_weeks"] >= 1
    assert isinstance(payload["weeks"], list)
    assert payload["weeks"]

    first_week = payload["weeks"][0]
    assert first_week["week_number"] == 1
    assert first_week["week_title"]
    assert first_week["total_hours"] >= 0
    assert isinstance(first_week["skills"], list)
    assert first_week["skills"]

    first_skill = first_week["skills"][0]
    assert first_skill["item_id"]
    assert first_skill["skill_name"]
    assert first_skill["skill_band"]
    assert isinstance(first_skill["subtopics"], list)
    return payload["plan_id"], first_skill["item_id"]


@pytest.mark.asyncio
async def test_learning_roadmap_requires_target_role(async_client: httpx.AsyncClient):
    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)

    response = await async_client.get("/api/v1/users/me/learning-roadmap", headers=headers)
    assert response.status_code == 400, response.text


@pytest.mark.asyncio
async def test_learning_roadmap_requires_processed_resume(async_client: httpx.AsyncClient):
    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)
    await _seed_target_role(suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    response = await async_client.get("/api/v1/users/me/learning-roadmap", headers=headers)
    assert response.status_code == 400, response.text


@pytest.mark.asyncio
async def test_learning_roadmap_full_flow(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)

    assert plan_id
    assert item_id


@pytest.mark.asyncio
async def test_learning_roadmap_replaces_unreachable_resource_urls(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    async def _mock_get_resources_for_skill(*, db, skill_name: str, level: str, limit_per_format: int = 1):
        _ = db
        _ = level
        _ = limit_per_format
        return {
            "video": [
                {
                    "title": f"{skill_name} Video",
                    "url": "https://invalid.example.com/video",
                    "provider": "Catalog",
                    "format": "video",
                    "duration_minutes": 45,
                }
            ],
            "article": [
                {
                    "title": f"{skill_name} Article",
                    "url": "https://invalid.example.com/article",
                    "provider": "Catalog",
                    "format": "article",
                    "duration_minutes": 30,
                }
            ],
        }

    async def _mock_is_url_reachable(client, url: str) -> bool:
        _ = client
        _ = url
        return False

    async def _mock_skill_rationale(
        skill_name: str,
        target_role: str,
        gap_type: str,
        rationale_cache: dict[tuple[str, str, str], str] | None = None,
    ) -> str:
        _ = skill_name
        _ = target_role
        _ = gap_type
        _ = rationale_cache
        return "Catalog rationale"

    monkeypatch.setattr(path_generator, "get_resources_for_skill", _mock_get_resources_for_skill)
    monkeypatch.setattr(path_generator, "is_url_reachable", _mock_is_url_reachable)
    monkeypatch.setattr(path_generator, "_gemini_skill_rationale", _mock_skill_rationale)

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)

    first_skill_resources = roadmap["weeks"][0]["skills"][0]["resources"]
    assert first_skill_resources
    assert all("invalid.example.com" not in str(resource.get("url", "")) for resource in first_skill_resources)
    assert all("google.com/search?q=" in str(resource.get("url", "")) for resource in first_skill_resources)

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        assert item.plan_id == plan_id

        _stored_subtopics, stored_resources = path_generator._deserialize_learning_content(item.subtopics_json)
        assert stored_resources
        assert all("invalid.example.com" not in resource.url for resource in stored_resources)
        assert all("google.com/search?q=" in resource.url for resource in stored_resources)

        normalized_rows = (
            await db.execute(
                select(LearningPlanItemResource)
                .where(LearningPlanItemResource.item_id == item_id)
                .order_by(LearningPlanItemResource.rank.asc())
            )
        ).scalars().all()
        assert normalized_rows
        assert all("invalid.example.com" not in str(row.url or "") for row in normalized_rows)
        assert all("google.com/search?q=" in str(row.url or "") for row in normalized_rows)


@pytest.mark.asyncio
async def test_learning_roadmap_accepts_daily_hours_without_40h_cap(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
):
    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    response = await async_client.get(
        "/api/v1/users/me/learning-roadmap?daily_hours=7&deadline_weeks=8",
        headers=headers,
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["hours_per_week"] == 49


@pytest.mark.asyncio
async def test_learning_roadmap_daily_hours_validation_errors(async_client: httpx.AsyncClient):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)

    response_non_positive = await async_client.get(
        "/api/v1/users/me/learning-roadmap?daily_hours=0&deadline_weeks=8",
        headers=headers,
    )
    assert response_non_positive.status_code == 422, response_non_positive.text
    assert "daily_hours" in response_non_positive.text

    response_too_high = await async_client.get(
        "/api/v1/users/me/learning-roadmap?daily_hours=25&deadline_weeks=8",
        headers=headers,
    )
    assert response_too_high.status_code == 422, response_too_high.text
    assert "daily_hours" in response_too_high.text


@pytest.mark.asyncio
async def test_learning_roadmap_idempotent(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    first = await _generate_roadmap(async_client, headers)
    second = await _generate_roadmap(async_client, headers)
    forced = await async_client.get(
        "/api/v1/users/me/learning-roadmap?hours_per_week=10&force_regenerate=true",
        headers=headers,
    )
    assert forced.status_code == 200, forced.text
    forced_payload = forced.json()

    assert first["plan_id"] == second["plan_id"]
    assert first["weeks"][0]["skills"][0]["item_id"] == second["weeks"][0]["skills"][0]["item_id"]
    assert forced_payload["plan_id"] != first["plan_id"]


@pytest.mark.asyncio
async def test_roadmap_respects_hours_per_week(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)

    async with db_session.AsyncSessionLocal() as db:
        role = Role(
            name=TARGET_ROLE_NAME,
            description=f"Hours budget role {suffix}",
            department=None,
            seniority_level="Senior",
            is_custom=True,
        )
        db.add(role)
        await db.flush()

        skill_small_one = Skill(
            name=f"Budget Skill A {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=2,
            time_to_learn_hours=2,
            source_role=role.name,
        )
        skill_small_two = Skill(
            name=f"Budget Skill B {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=2,
            time_to_learn_hours=3,
            source_role=role.name,
        )
        skill_oversized = Skill(
            name=f"Budget Skill C {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=4,
            time_to_learn_hours=8,
            source_role=role.name,
        )
        db.add_all([skill_small_one, skill_small_two, skill_oversized])
        await db.flush()

        db.add_all(
            [
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_small_one.id,
                    importance=1.0,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_small_two.id,
                    importance=0.9,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_oversized.id,
                    importance=0.8,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
            ]
        )
        await db.commit()

    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=5)
    assert roadmap["hours_per_week"] == 5

    skill_c_name = f"Budget Skill C {suffix}"
    skill_c_segments = 0
    saw_multi_skill_week = False
    for week in roadmap["weeks"]:
        week_hours = float(week["total_hours"])
        skills = week["skills"]
        skill_c_segments += sum(1 for skill in skills if skill["skill_name"] == skill_c_name)

        if len(skills) > 1:
            saw_multi_skill_week = True
            assert week_hours <= 5.0
        if week_hours > 5.0:
            # Allowed only when a single oversized subtopic segment occupies the week.
            assert len(skills) == 1
            assert len(skills[0].get("subtopics", [])) == 1
            assert float(skills[0]["subtopics"][0].get("estimated_hours") or 0.0) > 5.0

    assert saw_multi_skill_week
    assert skill_c_segments >= 2


@pytest.mark.asyncio
async def test_roadmap_respects_deadline_budget_with_overflow(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)

    async with db_session.AsyncSessionLocal() as db:
        role = Role(
            name=TARGET_ROLE_NAME,
            description=f"Deadline budget role {suffix}",
            department=None,
            seniority_level="Senior",
            is_custom=True,
        )
        db.add(role)
        await db.flush()

        skill_mandatory = Skill(
            name=f"Deadline Skill Mandatory {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=2,
            time_to_learn_hours=20,
            source_role=role.name,
        )
        skill_optional_a = Skill(
            name=f"Deadline Skill Optional A {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=2,
            time_to_learn_hours=12,
            source_role=role.name,
        )
        skill_optional_b = Skill(
            name=f"Deadline Skill Optional B {suffix}",
            skill_type="technical",
            category="Technical Skills",
            prerequisites="[]",
            difficulty=2,
            time_to_learn_hours=15,
            source_role=role.name,
        )
        db.add_all([skill_mandatory, skill_optional_a, skill_optional_b])
        await db.flush()

        db.add_all(
            [
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_mandatory.id,
                    importance=1.0,
                    is_mandatory=True,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_optional_a.id,
                    importance=0.7,
                    is_mandatory=False,
                    min_proficiency="intermediate",
                ),
                RoleSkillRequirement(
                    role_id=role.id,
                    skill_id=skill_optional_b.id,
                    importance=0.6,
                    is_mandatory=False,
                    min_proficiency="intermediate",
                ),
            ]
        )
        await db.commit()

    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(
        async_client,
        headers,
        hours_per_week=14,
        deadline_weeks=2,
        daily_hours=2,
    )

    budget_hours = 14 * 2
    assert roadmap["total_weeks"] <= 2
    assert float(roadmap["budget"]["scheduled_hours"]) <= float(budget_hours) + 1e-6
    assert len(roadmap.get("deferred_items", [])) == 0
    assert float(roadmap["budget"].get("overflow_hours_estimate") or 0) > 0
    scheduled_names = set()
    for week in roadmap.get("weeks") or []:
        for sk in week.get("skills") or []:
            scheduled_names.add(sk.get("skill_name"))
    assert len(scheduled_names) >= 3


@pytest.mark.asyncio
async def test_roadmap_large_budget_has_no_overflow(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(
        async_client,
        headers,
        hours_per_week=35,
        deadline_weeks=10,
        daily_hours=5,
    )

    assert roadmap.get("deferred_items", []) == []


@pytest.mark.asyncio
async def test_progress_reranks_remaining(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_prereq_chain_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap_before = await _generate_roadmap(async_client, headers)
    skill_positions_before: dict[str, int] = {}
    for week_index, week in enumerate(roadmap_before["weeks"]):
        for skill in week["skills"]:
            name = skill["skill_name"]
            skill_positions_before[name] = min(skill_positions_before.get(name, week_index), week_index)
    skill_ids = {
        skill["skill_name"]: skill["item_id"]
        for week in roadmap_before["weeks"]
        for skill in week["skills"]
    }

    completed_response = await async_client.patch(
        f"/api/v1/users/me/learning-plan/{roadmap_before['plan_id']}/items/{skill_ids[f'Chain Skill A {suffix}']}/progress",
        headers=headers,
        json={"status": "completed"},
    )
    assert completed_response.status_code == 200, completed_response.text
    completed_payload = completed_response.json()
    assert completed_payload["reranked"] is True

    roadmap_after = await _generate_roadmap(async_client, headers)
    skill_positions_after: dict[str, int] = {}
    for week_index, week in enumerate(roadmap_after["weeks"]):
        for skill in week["skills"]:
            name = skill["skill_name"]
            skill_positions_after[name] = min(skill_positions_after.get(name, week_index), week_index)

    skill_b = f"Chain Skill B {suffix}"
    skill_d = f"Chain Skill D {suffix}"

    assert skill_positions_before[skill_d] < skill_positions_before[skill_b]
    assert skill_positions_after[skill_b] < skill_positions_after[skill_d]


@pytest.mark.asyncio
async def test_progress_update_lifecycle(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)

    in_progress_response = await async_client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "in_progress"},
    )
    assert in_progress_response.status_code == 200, in_progress_response.text
    in_progress_payload = in_progress_response.json()
    assert in_progress_payload["status"] == "in_progress"
    assert in_progress_payload["plan_status"] == "in_progress"

    completed_response = await async_client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "completed"},
    )
    assert completed_response.status_code == 200, completed_response.text
    completed_payload = completed_response.json()
    assert completed_payload["completed_at"] is not None
    assert completed_payload["plan_status"] in {"in_progress", "completed"}

    revert_response = await async_client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "not_started"},
    )
    assert revert_response.status_code == 400, revert_response.text


@pytest.mark.asyncio
async def test_progress_summary(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)

    completed_response = await async_client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "completed"},
    )
    assert completed_response.status_code == 200, completed_response.text

    summary_response = await async_client.get("/api/v1/users/me/learning-plan/progress", headers=headers)
    assert summary_response.status_code == 200, summary_response.text
    summary = summary_response.json()

    assert summary["plan_id"] == plan_id
    assert summary["completed_items"] >= 1
    assert summary["percent_complete"] > 0
    assert summary["hours_completed"] > 0
    assert summary["total_items"] == summary["completed_items"] + summary["in_progress_items"] + summary["not_started_items"]


@pytest.mark.asyncio
async def test_progress_summary_no_plan(async_client: httpx.AsyncClient):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)

    response = await async_client.get("/api/v1/users/me/learning-plan/progress", headers=headers)
    assert response.status_code == 404, response.text


@pytest.mark.asyncio
async def test_assessment_feedback_low_score(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)
    skill_name = roadmap["weeks"][0]["skills"][0]["skill_name"]
    original_subtopics = roadmap["weeks"][0]["skills"][0]["subtopics"]
    original_resources = roadmap["weeks"][0]["skills"][0]["resources"]

    response = await async_client.post(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/assessment-feedback",
        headers=headers,
        json={
            "skill_name": skill_name,
            "score": 0.4,
            "failed_areas": ["api_design", "error_handling"],
            "passed_areas": ["basics"],
        },
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["action_taken"] == "subtopics_regenerated"
    assert payload["reranked_remaining"] is False
    assert payload["updated_subtopics"]

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        assert item.status == "needs_review"
        assert item.subtopics_json is not None
        stored_subtopics, stored_resources = path_generator._deserialize_learning_content(item.subtopics_json)
        assert stored_subtopics
        assert len(stored_resources) == len(original_resources)


@pytest.mark.asyncio
async def test_assessment_feedback_high_score(async_client: httpx.AsyncClient, tmp_path: Path):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)
    skill_name = roadmap["weeks"][0]["skills"][0]["skill_name"]

    response = await async_client.post(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/assessment-feedback",
        headers=headers,
        json={
            "skill_name": skill_name,
            "score": 0.8,
            "failed_areas": ["api_design"],
            "passed_areas": ["basics", "implementation"],
        },
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["action_taken"] == "upgraded_to_weak"
    assert payload["updated_subtopics"] is None
    assert payload["reranked_remaining"] is False

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        assert item.resource_type == "weak"


@pytest.mark.asyncio
async def test_item_assessment_get_submit_and_history(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(assessments_endpoint, "generate_assessment_questions", _mock_generate_assessment_questions)

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)
    _ = plan_id

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        item.assessment_questions = None
        db.add(item)
        await db.commit()

    get_resp = await async_client.get(f"/api/v1/assessments/items/{item_id}", headers=headers)
    assert get_resp.status_code == 200, get_resp.text
    get_payload = get_resp.json()
    assert get_payload["attempts"] == 0
    assert get_payload["last_score"] is None
    assert len(get_payload["questions"]) == 5
    assert set(get_payload["questions"][0].keys()) == {"index", "question", "options"}

    submit_resp = await async_client.post(
        f"/api/v1/assessments/items/{item_id}/submit",
        headers=headers,
        json={"answers": [0, 0, 2, 0, 0]},
    )
    assert submit_resp.status_code == 200, submit_resp.text
    submit_payload = submit_resp.json()
    assert submit_payload["passed"] is False
    assert submit_payload["correct_count"] == 3
    assert submit_payload["total"] == 5
    assert submit_payload["item_status"] == "needs_review"

    history_resp = await async_client.get(f"/api/v1/assessments/items/{item_id}/history", headers=headers)
    assert history_resp.status_code == 200, history_resp.text
    history_payload = history_resp.json()
    assert history_payload["item_id"] == item_id
    assert len(history_payload["attempts"]) == 1
    assert history_payload["attempts"][0]["max_score"] == 5
    assert history_payload["attempts"][0]["score"] == 3


@pytest.mark.asyncio
async def test_item_assessment_pass_returns_remediation_and_schedules_next_week_generation(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(assessments_endpoint, "generate_assessment_questions", _mock_generate_assessment_questions)

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        item.assessment_questions = None
        db.add(item)
        await db.commit()

    get_resp = await async_client.get(f"/api/v1/assessments/items/{item_id}", headers=headers)
    assert get_resp.status_code == 200, get_resp.text

    next_week_item_ids = ["week-2-item-a", "week-2-item-b"]

    async def _mock_resolve_item_week_context(db, plan, item):
        _ = db
        _ = plan
        _ = item
        return 1, [item_id], next_week_item_ids

    async def _mock_is_week_complete(db, plan_id, week_item_ids, status_overrides=None):
        _ = db
        _ = plan_id
        _ = week_item_ids
        _ = status_overrides
        return True

    schedule_calls = {"count": 0, "args": None}

    async def _mock_generate_week_item_questions_background(plan_id_arg, item_ids_arg):
        schedule_calls["count"] += 1
        schedule_calls["args"] = (plan_id_arg, list(item_ids_arg))

    monkeypatch.setattr(assessments_endpoint, "_resolve_item_week_context", _mock_resolve_item_week_context)
    monkeypatch.setattr(assessments_endpoint, "_is_week_complete", _mock_is_week_complete)
    monkeypatch.setattr(
        assessments_endpoint,
        "_generate_week_item_questions_background",
        _mock_generate_week_item_questions_background,
    )

    # Mock questions have correct indexes [0,1,2,3,0]; final answer is intentionally wrong => 4/5 pass.
    submit_resp = await async_client.post(
        f"/api/v1/assessments/items/{item_id}/submit",
        headers=headers,
        json={"answers": [0, 1, 2, 3, 1]},
    )
    assert submit_resp.status_code == 200, submit_resp.text
    payload = submit_resp.json()

    assert payload["passed"] is True
    assert payload["correct_count"] == 4
    assert payload["item_status"] == "completed"
    assert "results" in payload
    assert isinstance(payload["results"], list)
    assert len(payload["results"]) == 5
    assert payload["results"][0]["selected_index"] == 0
    assert payload["results"][0]["selected_option"] == "A"
    assert payload["results"][0]["correct_option"] == "A"
    assert payload["results"][4]["selected_index"] == 1
    assert payload["results"][4]["correct_index"] == 0
    assert payload["results"][4]["selected_option"] == "B"
    assert payload["results"][4]["correct_option"] == "A"
    assert "remediation_feedback" in payload
    assert isinstance(payload["remediation_feedback"], list)
    assert len(payload["remediation_feedback"]) == 1
    feedback = payload["remediation_feedback"][0]
    assert feedback["question_index"] == 4
    assert feedback["user_answer"] == 1
    assert feedback["correct_index"] == 0
    assert isinstance(feedback["explanation"], str)
    assert isinstance(feedback["quick_tip"], str)
    assert feedback["quick_tip"]

    assert schedule_calls["count"] == 1
    assert schedule_calls["args"] == (plan_id, next_week_item_ids)

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        assert item.status == "completed"


@pytest.mark.asyncio
async def test_item_assessment_get_uses_cached_questions_without_generation(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    _, item_id = _assert_roadmap_payload(roadmap)

    cached_questions = await _mock_generate_assessment_questions("cached", "beginner", [])
    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        item.assessment_questions = json.dumps(cached_questions)
        db.add(item)
        await db.commit()

    async def _should_not_be_called(skill_name: str, level: str, subtopics: list[str]) -> list[dict]:
        _ = skill_name
        _ = level
        _ = subtopics
        raise AssertionError("generate_assessment_questions should not be called for cached items")

    monkeypatch.setattr(assessments_endpoint, "generate_assessment_questions", _should_not_be_called)

    response = await async_client.get(f"/api/v1/assessments/items/{item_id}", headers=headers)
    assert response.status_code == 200, response.text
    payload = response.json()
    assert len(payload["questions"]) == 5
    assert set(payload["questions"][0].keys()) == {"index", "question", "options"}


@pytest.mark.asyncio
async def test_item_assessment_get_generates_once_and_reuses_persisted_questions(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    _, item_id = _assert_roadmap_payload(roadmap)

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        item.assessment_questions = None
        db.add(item)
        await db.commit()

    calls = {"count": 0}

    async def _counting_generator(skill_name: str, level: str, subtopics: list[str]) -> list[dict]:
        _ = skill_name
        _ = level
        _ = subtopics
        calls["count"] += 1
        return await _mock_generate_assessment_questions(skill_name, level, subtopics)

    monkeypatch.setattr(assessments_endpoint, "generate_assessment_questions", _counting_generator)

    first = await async_client.get(f"/api/v1/assessments/items/{item_id}", headers=headers)
    second = await async_client.get(f"/api/v1/assessments/items/{item_id}", headers=headers)
    assert first.status_code == 200, first.text
    assert second.status_code == 200, second.text
    assert calls["count"] == 1

    async with db_session.AsyncSessionLocal() as db:
        item = await db.scalar(select(LearningPlanItem).where(LearningPlanItem.id == item_id))
        assert item is not None
        parsed = json.loads(item.assessment_questions or "[]")
        assert isinstance(parsed, list)
        assert len(parsed) == 5


@pytest.mark.asyncio
async def test_manual_completion_blocked_when_assessment_exists(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(assessments_endpoint, "generate_assessment_questions", _mock_generate_assessment_questions)

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, item_id = _assert_roadmap_payload(roadmap)

    get_resp = await async_client.get(f"/api/v1/assessments/items/{item_id}", headers=headers)
    assert get_resp.status_code == 200, get_resp.text

    progress_resp = await async_client.patch(
        f"/api/v1/users/me/learning-plan/{plan_id}/items/{item_id}/progress",
        headers=headers,
        json={"status": "completed"},
    )
    assert progress_resp.status_code == 400, progress_resp.text
    assert "requires assessment submission" in progress_resp.text.lower()


@pytest.mark.asyncio
async def test_duplicate_skills_generate_item_unique_assessments(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    async def _mock_detect_duplicate_gaps(db: AsyncSession, user, role):
        req = await db.scalar(
            select(RoleSkillRequirement)
            .where(RoleSkillRequirement.role_id == role.id)
            .order_by(RoleSkillRequirement.importance.desc())
            .limit(1)
        )
        assert req is not None
        skill = await db.scalar(select(Skill).where(Skill.id == req.skill_id))
        assert skill is not None

        duplicate_gap_one = GapItem(
            skill_id=skill.id,
            skill_name=skill.name,
            gap_type="missing",
            priority_score=0.9,
            current_proficiency=None,
            required_proficiency="intermediate",
            time_to_learn_hours=int(skill.time_to_learn_hours or 10),
            importance=1.0,
            is_mandatory=True,
            prerequisites=[],
            prerequisite_coverage=1.0,
            skill_band="Technical Skills",
        )
        duplicate_gap_two = GapItem(
            skill_id=skill.id,
            skill_name=skill.name,
            gap_type="missing",
            priority_score=0.85,
            current_proficiency=None,
            required_proficiency="intermediate",
            time_to_learn_hours=int(skill.time_to_learn_hours or 10),
            importance=1.0,
            is_mandatory=True,
            prerequisites=[],
            prerequisite_coverage=1.0,
            skill_band="Technical Skills",
        )
        return GapResult(
            user_id=user.id,
            role_id=role.id,
            role_name=role.name,
            readiness_score=0.0,
            missing_skills=2,
            weak_skills=0,
            total_gaps=2,
            total_learning_hours=duplicate_gap_one.time_to_learn_hours * 2,
            gaps=[duplicate_gap_one, duplicate_gap_two],
        )

    call_counter = {"n": 0}

    async def _mock_item_question_generation(skill_name: str, level: str, subtopics: list[str]) -> list[dict]:
        _ = skill_name
        _ = level
        _ = subtopics
        call_counter["n"] += 1
        marker = call_counter["n"]
        return [
            {
                "question": f"Question {i + 1} marker {marker}",
                "options": ["A", "B", "C", "D"],
                "correct_index": 0,
                "explanation": f"Explanation marker {marker}",
            }
            for i in range(5)
        ]

    monkeypatch.setattr(path_generator, "detect_and_store_gaps", _mock_detect_duplicate_gaps)
    monkeypatch.setattr(path_generator, "generate_assessment_questions", _mock_item_question_generation)

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers)
    plan_id, _item_id = _assert_roadmap_payload(roadmap)

    async with db_session.AsyncSessionLocal() as db:
        items = (
            await db.execute(
                select(LearningPlanItem)
                .where(LearningPlanItem.plan_id == plan_id)
                .order_by(LearningPlanItem.order.asc())
            )
        ).scalars().all()

        assert len(items) == 2
        assert all(item.assessment_questions is not None for item in items)
        assert items[0].assessment_questions != items[1].assessment_questions
        assert "marker 1" in items[0].assessment_questions
        assert "marker 2" in items[1].assessment_questions
        assert call_counter["n"] == 2


@pytest.mark.asyncio
async def test_week_assessment_get_dynamic_count_and_tags(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        week_assessments_endpoint,
        "generate_week_assessment_questions",
        _mock_generate_week_assessment_questions,
    )

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=80)
    assert len(roadmap["weeks"]) >= 2

    plan_id = roadmap["plan_id"]
    response = await async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers)
    assert response.status_code == 200, response.text
    payload = response.json()

    assert payload["week_number"] == 1
    assert 8 <= int(payload["question_count"]) <= 20
    assert len(payload["questions"]) == int(payload["question_count"])
    assert payload["questions"]
    first_question = payload["questions"][0]
    assert "tags" in first_question
    assert "correct_index" not in first_question
    assert "explanation" not in first_question


@pytest.mark.asyncio
async def test_week_assessment_submit_pass_marks_completed_and_schedules_next_week(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        week_assessments_endpoint,
        "generate_week_assessment_questions",
        _mock_generate_week_assessment_questions,
    )

    scheduled = {"calls": []}

    async def _mock_background(user_id: str, plan_id: str, week_number: int) -> None:
        scheduled["calls"].append((user_id, plan_id, week_number))

    monkeypatch.setattr(
        week_assessments_endpoint,
        "ensure_week_assessment_generated_background",
        _mock_background,
    )

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=80)
    assert len(roadmap["weeks"]) >= 2
    plan_id = roadmap["plan_id"]

    get_resp = await async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers)
    assert get_resp.status_code == 200, get_resp.text
    question_count = int(get_resp.json()["question_count"])
    answers = [idx % 4 for idx in range(question_count)]

    submit_resp = await async_client.post(
        f"/api/v1/assessments/weeks/{plan_id}/1/submit",
        headers=headers,
        json={"answers": answers},
    )
    assert submit_resp.status_code == 200, submit_resp.text
    payload = submit_resp.json()
    assert payload["passed"] is True
    assert payload["status"] == "completed"
    assert payload["correct_count"] == payload["total"]

    assert len(scheduled["calls"]) == 1
    _user_id, scheduled_plan_id, scheduled_week = scheduled["calls"][0]
    assert scheduled_plan_id == plan_id
    assert scheduled_week == 2


@pytest.mark.asyncio
async def test_week_assessment_submit_after_pass_is_blocked_and_does_not_create_attempt(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        week_assessments_endpoint,
        "generate_week_assessment_questions",
        _mock_generate_week_assessment_questions,
    )

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=80)
    assert len(roadmap["weeks"]) >= 2
    plan_id = roadmap["plan_id"]

    get_resp = await async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers)
    assert get_resp.status_code == 200, get_resp.text
    question_count = int(get_resp.json()["question_count"])
    answers = [idx % 4 for idx in range(question_count)]

    submit_resp = await async_client.post(
        f"/api/v1/assessments/weeks/{plan_id}/1/submit",
        headers=headers,
        json={"answers": answers},
    )
    assert submit_resp.status_code == 200, submit_resp.text

    async with db_session.AsyncSessionLocal() as db:
        row = await db.scalar(
            select(WeekAssessment).where(
                WeekAssessment.plan_id == plan_id,
                WeekAssessment.week_number == 1,
            )
        )
        assert row is not None
        assert row.status == "completed"

        after_first_count = int(
            await db.scalar(
                select(func.count(WeekAssessmentAttempt.id)).where(
                    WeekAssessmentAttempt.week_assessment_id == row.id,
                    WeekAssessmentAttempt.user_id == row.user_id,
                )
            )
            or 0
        )
        assert after_first_count == 1

    second_resp = await async_client.post(
        f"/api/v1/assessments/weeks/{plan_id}/1/submit",
        headers=headers,
        json={"answers": answers},
    )
    assert second_resp.status_code == 409, second_resp.text
    assert "retakes are not allowed" in second_resp.text.lower()

    async with db_session.AsyncSessionLocal() as db:
        row = await db.scalar(
            select(WeekAssessment).where(
                WeekAssessment.plan_id == plan_id,
                WeekAssessment.week_number == 1,
            )
        )
        assert row is not None
        assert row.status == "completed"
        after_second_count = int(
            await db.scalar(
                select(func.count(WeekAssessmentAttempt.id)).where(
                    WeekAssessmentAttempt.week_assessment_id == row.id,
                    WeekAssessmentAttempt.user_id == row.user_id,
                )
            )
            or 0
        )
        assert after_second_count == after_first_count


@pytest.mark.asyncio
async def test_week_assessment_submit_wrong_length_answers_returns_400(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        week_assessments_endpoint,
        "generate_week_assessment_questions",
        _mock_generate_week_assessment_questions,
    )

    suffix = uuid.uuid4().hex[:8]
    headers = await _create_fresh_user(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=80)
    plan_id = roadmap["plan_id"]

    get_resp = await async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers)
    assert get_resp.status_code == 200, get_resp.text
    question_count = int(get_resp.json()["question_count"])
    assert question_count >= 2

    bad_answers = [0] * (question_count - 1)
    submit_resp = await async_client.post(
        f"/api/v1/assessments/weeks/{plan_id}/1/submit",
        headers=headers,
        json={"answers": bad_answers},
    )
    assert submit_resp.status_code == 400, submit_resp.text


@pytest.mark.asyncio
async def test_week_assessment_distribution_enforced_with_retry(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        week_assessments_endpoint,
        "generate_week_assessment_questions",
        week_assessment_service.generate_week_assessment_questions,
    )

    calls = {"count": 0}

    def _build_llm_payload(question_count: int, edge_count: int, scenario_count: int) -> str:
        payload = []
        for idx in range(question_count):
            tags = ["skill:Mock", "subtopic:Mock"]
            if idx < edge_count:
                tags.append("type:edge_case")
            elif idx < edge_count + scenario_count:
                tags.append("type:scenario")
            else:
                tags.append("type:concept")
            payload.append(
                {
                    "question": f"Q{idx + 1}",
                    "options": ["A", "B", "C", "D"],
                    "correct_index": idx % 4,
                    "explanation": f"E{idx + 1}",
                    "tags": tags,
                }
            )
        return json.dumps(payload)

    async def _mock_llm_generate(*, purpose: str, prompt: str, user_id: str | None = None, request_id: str | None = None):
        _ = purpose
        _ = user_id
        _ = request_id
        calls["count"] += 1

        q_match = re.search(r"exactly\s+(\d+)\s+objects", prompt)
        e_match = re.search(r"At least\s+(\d+)\s+questions tagged as edge", prompt)
        s_match = re.search(r"At least\s+(\d+)\s+scenario-based", prompt)
        question_count = int(q_match.group(1)) if q_match else 8
        min_edge = int(e_match.group(1)) if e_match else 2
        min_scenario = int(s_match.group(1)) if s_match else 3

        if calls["count"] == 1:
            # Intentionally fail edge requirement to force retry.
            return _build_llm_payload(question_count, edge_count=max(0, min_edge - 1), scenario_count=min_scenario)
        return _build_llm_payload(question_count, edge_count=min_edge, scenario_count=min_scenario)

    monkeypatch.setattr(week_assessment_service, "gemini_generate", _mock_llm_generate)

    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=80)
    plan_id = roadmap["plan_id"]

    # Clear bootstrap-generated questions so endpoint must regenerate with retry logic.
    async with db_session.AsyncSessionLocal() as db:
        row = await db.scalar(
            select(WeekAssessment).where(
                WeekAssessment.user_id == user_id,
                WeekAssessment.plan_id == plan_id,
                WeekAssessment.week_number == 1,
            )
        )
        assert row is not None
        row.questions_json = None
        row.status = "pending"
        db.add(row)
        await db.commit()

    response = await async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers)
    assert response.status_code == 200, response.text
    payload = response.json()
    assert len(payload["questions"]) == int(payload["question_count"])
    assert calls["count"] >= 2

    async with db_session.AsyncSessionLocal() as db:
        row = await db.scalar(
            select(WeekAssessment).where(
                WeekAssessment.user_id == user_id,
                WeekAssessment.plan_id == plan_id,
                WeekAssessment.week_number == 1,
            )
        )
        assert row is not None
        full_questions = json.loads(row.questions_json or "[]")
        assert len(full_questions) == int(row.question_count)

    edge_count = sum(1 for q in full_questions if any("edge" in str(t).lower() for t in q.get("tags", [])))
    scenario_count = sum(1 for q in full_questions if any("scenario" in str(t).lower() for t in q.get("tags", [])))
    assert edge_count >= max(2, int(math.ceil(int(row.question_count) * 0.2)))
    assert scenario_count >= int(math.ceil(int(row.question_count) * 0.3))


@pytest.mark.asyncio
async def test_week_assessment_user_isolation_allows_same_plan_week_across_users(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
):
    suffix_owner = uuid.uuid4().hex[:8]
    owner_user_id, _owner_email, owner_headers = await _create_user_and_token(async_client, suffix_owner)
    await _seed_target_role(suffix_owner)
    await _upload_and_process_resume(async_client, owner_headers, tmp_path, suffix_owner)
    await _set_target_role(async_client, owner_headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, owner_headers, hours_per_week=80)
    plan_id = roadmap["plan_id"]

    suffix_other = uuid.uuid4().hex[:8]
    other_user_id, _other_email, _other_headers = await _create_user_and_token(async_client, suffix_other)

    rogue_questions = json.dumps(
        [
            {
                "question": "ROGUE QUESTION",
                "options": ["A", "B", "C", "D"],
                "correct_index": 0,
                "explanation": "ROGUE EXPLANATION",
                "tags": ["skill:Rogue", "type:edge_case"],
            }
        ]
        * 8
    )

    async with db_session.AsyncSessionLocal() as db:
        db.add(
            WeekAssessment(
                user_id=other_user_id,
                plan_id=plan_id,
                week_number=1,
                question_count=8,
                questions_json=rogue_questions,
                status="ready",
            )
        )
        await db.commit()

    response = await async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=owner_headers)
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["questions"]
    assert payload["questions"][0]["question"] != "ROGUE QUESTION"

    async with db_session.AsyncSessionLocal() as db:
        rows = (
            await db.execute(
                select(WeekAssessment).where(
                    WeekAssessment.plan_id == plan_id,
                    WeekAssessment.week_number == 1,
                )
            )
        ).scalars().all()
        assert len(rows) == 2
        assert {row.user_id for row in rows} == {owner_user_id, other_user_id}


@pytest.mark.asyncio
async def test_week_assessment_concurrent_get_generation_is_idempotent(
    async_client: httpx.AsyncClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    markers = {"count": 0}

    async def _slow_generator(*, week_number: int, skills: list[str], subtopics: list[str], question_count: int) -> list[dict]:
        _ = week_number
        _ = skills
        _ = subtopics
        markers["count"] += 1
        marker = markers["count"]
        await asyncio.sleep(0.05)
        return [
            {
                "question": f"Race question marker {marker} #{idx + 1}",
                "options": ["A", "B", "C", "D"],
                "correct_index": idx % 4,
                "explanation": f"Race explanation {idx + 1}",
                "tags": ["skill:Race", "type:scenario"],
            }
            for idx in range(question_count)
        ]

    monkeypatch.setattr(week_assessments_endpoint, "generate_week_assessment_questions", _slow_generator)

    suffix = uuid.uuid4().hex[:8]
    user_id, _email, headers = await _create_user_and_token(async_client, suffix)
    await _seed_target_role(suffix)
    await _upload_and_process_resume(async_client, headers, tmp_path, suffix)
    await _set_target_role(async_client, headers, TARGET_ROLE_NAME)

    roadmap = await _generate_roadmap(async_client, headers, hours_per_week=80)
    plan_id = roadmap["plan_id"]

    async with db_session.AsyncSessionLocal() as db:
        row = await db.scalar(
            select(WeekAssessment).where(
                WeekAssessment.user_id == user_id,
                WeekAssessment.plan_id == plan_id,
                WeekAssessment.week_number == 1,
            )
        )
        assert row is not None
        row.questions_json = None
        row.status = "pending"
        db.add(row)
        await db.commit()

    first_resp, second_resp = await asyncio.gather(
        async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers),
        async_client.get(f"/api/v1/assessments/weeks/{plan_id}/1", headers=headers),
    )
    assert first_resp.status_code == 200, first_resp.text
    assert second_resp.status_code == 200, second_resp.text

    first_payload = first_resp.json()
    second_payload = second_resp.json()
    assert first_payload["questions"][0]["question"] == second_payload["questions"][0]["question"]

    async with db_session.AsyncSessionLocal() as db:
        row = await db.scalar(
            select(WeekAssessment).where(
                WeekAssessment.user_id == user_id,
                WeekAssessment.plan_id == plan_id,
                WeekAssessment.week_number == 1,
            )
        )
        assert row is not None
        stored_questions = json.loads(row.questions_json or "[]")
        assert stored_questions
        assert stored_questions[0]["question"] == first_payload["questions"][0]["question"]
